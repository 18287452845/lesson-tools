"""Debug script to analyze template vs output structure."""
from docx import Document
from pathlib import Path
import re
import sys

def analyze_document(filepath, name):
    """Analyze a Word document structure."""
    print(f"\n{'='*80}")
    print(f"Analyzing: {name}")
    print(f"File: {filepath}")
    print('='*80)
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"ERROR: Could not load document: {e}")
        return
    
    # Analyze overall structure
    print(f"\n📊 DOCUMENT STRUCTURE:")
    print(f"   Total Paragraphs: {len(doc.paragraphs)}")
    print(f"   Total Tables: {len(doc.tables)}")
    
    # Analyze tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📋 Table {table_idx}:")
        print(f"   Size: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Show structure with placeholders
        pattern = re.compile(r'\{\{[^}]+\}\}')
        print(f"\n   First 8 rows content:")
        for row_idx, row in enumerate(table.rows[:8]):
            cells_text = []
            all_placeholders = []
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip().replace('\n', ' ')[:60]
                if text:
                    cells_text.append(f"Col{col_idx}:{text}")
                placeholders = pattern.findall(cell.text)
                if placeholders:
                    all_placeholders.extend(placeholders)
            
            if cells_text or all_placeholders:
                print(f"   Row {row_idx:2d}: {' | '.join(cells_text)}")
                if all_placeholders:
                    print(f"         → {', '.join(all_placeholders)}")
    
    # Analyze paragraphs
    print(f"\n📝 PARAGRAPHS (first 10 non-empty):")
    count = 0
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # Look for placeholders
            placeholders = re.findall(r'\{\{[^}]+\}\}', text)
            placeholder_text = f" → {', '.join(placeholders)}" if placeholders else ""
            print(f"   [{idx:2d}] {text[:70]}{placeholder_text}")
            count += 1
            if count >= 10:
                break
    
    print('\n' + '='*80 + '\n')

if __name__ == "__main__":
    # Analyze template
    template_path = "storage/templates/yunnan_forestry_college_template.docx"
    print("\n🔍 COMPARING TEMPLATE vs OUTPUT\n")
    
    if Path(template_path).exists():
        analyze_document(template_path, "📄 TEMPLATE")
    else:
        print(f"ERROR: Template not found at {template_path}")
        sys.exit(1)
    
    # Analyze latest output
    outputs = sorted(Path("storage/outputs").glob("*.docx"), 
                     key=lambda x: x.stat().st_mtime, reverse=True)
    if outputs:
        analyze_document(outputs[0], "📄 LATEST OUTPUT")
        
        # Show differences
        print("\n⚠️  STRUCTURE COMPARISON:")
        template_doc = Document(template_path)
        output_doc = Document(outputs[0])
        
        print(f"   Paragraphs: Template={len(template_doc.paragraphs)}, Output={len(output_doc.paragraphs)}")
        print(f"   Tables: Template={len(template_doc.tables)}, Output={len(output_doc.tables)}")
        
        for i in range(min(len(template_doc.tables), len(output_doc.tables))):
            t_table = template_doc.tables[i]
            o_table = output_doc.tables[i]
            t_size = f"{len(t_table.rows)}x{len(t_table.columns)}"
            o_size = f"{len(o_table.rows)}x{len(o_table.columns)}"
            match = "✅" if t_size == o_size else "❌"
            print(f"   Table {i}: Template={t_size}, Output={o_size} {match}")
    else:
        print("No output files found.")

