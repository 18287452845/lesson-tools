"""
Build competition Word templates from scratch using python-docx + docxtpl Jinja2 syntax.

Generates two templates in storage/competition_templates/:
  - 参赛教案_模板.docx
  - 教学实施报告_模板.docx

Run from project root:
    python build_competition_templates.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = Path(__file__).parent / "storage" / "competition_templates"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_background(cell, color_hex: str) -> None:
    """Set table cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_centered_heading(doc, text: str, font_size: int = 24, bold: bool = True) -> None:
    """Add a centered heading paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_paragraph(doc, text: str, font_size: int = 11, bold: bool = False, indent_first: bool = True) -> None:
    """Add a regular paragraph."""
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_section_title(doc, text: str, level: int = 2) -> None:
    """Add a section title (一、 / (一) etc)."""
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_label_value_table(doc, rows: list[tuple[str, str]]) -> None:
    """Add a label-value table for cover info."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value_template) in enumerate(rows):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = value_template
        # Style label
        for p in cell_label.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)
        cell_label.width = Cm(4)
        cell_value.width = Cm(11)


def _add_steps_table(doc, list_path: str, include_political: bool = False) -> None:
    """
    Add a teaching-steps table using the 4-row pattern:
      Row 0: header
      Row 1: {%tr for step in <list_path> %}    (whole row is a tr-loop OPEN marker)
      Row 2: data row with {{ step.* }} variables
      Row 3: {%tr endfor %}                     (whole row is a tr-loop CLOSE marker)

    docxtpl will:
      - delete row 1 and row 3 (they only contain control tags)
      - replicate row 2 for each item in the list
    """
    headers = ["教学环节", "教学内容", "教师活动", "学生活动", "设计意图"]
    if include_political:
        headers.append("思政点")
    cols = len(headers)

    table = doc.add_table(rows=4, cols=cols)
    table.style = "Table Grid"

    # Row 0: headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_background(cell, "DEEAF6")

    # Row 1: {%tr for ... %} marker (occupies first cell only)
    table.rows[1].cells[0].text = "{%tr for step in " + list_path + " %}"

    # Row 2: data placeholders
    data_cells = table.rows[2].cells
    if include_political:
        data_cells[0].text = "{{ step.stage }}{% if step.duration %}({{ step.duration }}){% endif %}"
        data_cells[1].text = "{{ step.content }}"
        data_cells[2].text = "{{ step.teacher_activity }}"
        data_cells[3].text = "{{ step.student_activity }}"
        data_cells[4].text = "{{ step.design_intent }}"
        data_cells[5].text = "{{ step.ideological_political }}"
    else:
        data_cells[0].text = "{{ step.stage }}"
        data_cells[1].text = "{{ step.content }}"
        data_cells[2].text = "{{ step.teacher_activity }}"
        data_cells[3].text = "{{ step.student_activity }}"
        data_cells[4].text = "{{ step.design_intent }}"

    # Row 3: {%tr endfor %} marker
    table.rows[3].cells[0].text = "{%tr endfor %}"


# ============================================================================
# Template 1: 参赛教案_模板.docx
# ============================================================================

def build_lesson_plan_template() -> None:
    doc = Document()

    # ----- Cover -----
    for _ in range(3):
        doc.add_paragraph()
    add_centered_heading(doc, "{{ competition_year }}{{ competition_region }}职业院校技能大赛", font_size=20)
    add_centered_heading(doc, "教学能力比赛", font_size=22)
    for _ in range(2):
        doc.add_paragraph()
    add_centered_heading(doc, "参 赛 教 案", font_size=36)
    for _ in range(3):
        doc.add_paragraph()

    add_label_value_table(doc, [
        ("作品名称：", "{{ work_name }}"),
        ("课程名称：", "{{ course_name }}"),
        ("专业大类：", "{{ major_category }}"),
        ("参赛组别：", "{{ group_name }}"),
        ("授课班级：", "{{ class_name }}"),
        ("总学时：", "{{ total_hours }} 学时"),
    ])
    doc.add_page_break()

    # ----- 一、整体设计 -----
    add_section_title(doc, "一、整体设计", level=1)
    doc.add_paragraph()

    add_section_title(doc, "(一) 内容分析", level=2)
    add_paragraph(doc, "{{ overall_design.content_analysis }}")
    doc.add_paragraph()

    add_section_title(doc, "(二) 学情分析", level=2)
    add_paragraph(doc, "{{ overall_design.student_analysis }}")
    doc.add_paragraph()

    add_section_title(doc, "(三) 目标分析", level=2)
    add_paragraph(doc, "{{ overall_design.goal_analysis }}")
    doc.add_paragraph()

    add_section_title(doc, "(四) 过程设计", level=2)
    add_paragraph(doc, "{{ overall_design.process_design }}")
    doc.add_paragraph()

    add_section_title(doc, "(五) 教学方法", level=2)
    add_paragraph(doc, "{{ overall_design.teaching_method }}")
    doc.add_paragraph()

    # 学情调查问卷
    add_section_title(doc, "附件 1 学情调查问卷", level=3)
    p = doc.add_paragraph()
    p.add_run("{%p for q in overall_design.survey_questions %}")
    add_paragraph(doc, "{{ loop.index }}、{{ q }}", indent_first=False)
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")
    doc.add_page_break()

    # ----- 二、任务详细设计 (loop over lessons) -----
    add_section_title(doc, "二、任务详细设计", level=1)

    # Loop opener at paragraph level
    p = doc.add_paragraph()
    p.add_run("{%p for lesson in lessons %}")

    # Per-lesson title
    add_section_title(doc, "【教案{{ lesson.lesson_number }}】{{ lesson.title }}", level=2)
    doc.add_paragraph()

    # 一、教学基本信息
    add_section_title(doc, "一、教学基本信息", level=3)
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    cells = info_table.rows[0].cells
    cells[0].text = "模块名称"
    cells[1].text = "{{ lesson.module_name }}"
    cells[2].text = "项目名称"
    cells[3].text = "{{ lesson.project_name }}"
    cells = info_table.rows[1].cells
    cells[0].text = "任务名称"
    cells[1].text = "{{ lesson.title }}"
    cells[2].text = "授课学时"
    cells[3].text = "{{ lesson.hours }}"
    cells = info_table.rows[2].cells
    cells[0].text = "授课地点"
    cells[1].text = "{{ lesson.location }}"
    cells[2].text = "授课班级"
    cells[3].text = "{{ lesson.class_name }}"
    doc.add_paragraph()

    # 二、教学设计
    add_section_title(doc, "二、教学设计", level=3)

    add_paragraph(doc, "【岗课赛证融合设计】", bold=True, indent_first=False)
    add_paragraph(doc, "{{ lesson.position_competition_certificate }}")
    doc.add_paragraph()

    add_paragraph(doc, "【学情分析】", bold=True, indent_first=False)
    add_paragraph(doc, "知识与技能基础:{{ lesson.student_analysis.knowledge_basis }}")
    add_paragraph(doc, "认知和实践能力:{{ lesson.student_analysis.cognition_practice }}")
    add_paragraph(doc, "学习特点:{{ lesson.student_analysis.learning_features }}")
    add_paragraph(doc, "评估结果:{{ lesson.student_analysis.assessment }}")
    doc.add_paragraph()

    add_paragraph(doc, "【教学目标】", bold=True, indent_first=False)
    add_paragraph(doc, "知识目标:", bold=True)
    p = doc.add_paragraph()
    p.add_run("{%p for k in lesson.objectives.knowledge %}")
    add_paragraph(doc, "{{ loop.index }}. {{ k }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    add_paragraph(doc, "能力目标:", bold=True)
    p = doc.add_paragraph()
    p.add_run("{%p for a in lesson.objectives.ability %}")
    add_paragraph(doc, "{{ loop.index }}. {{ a }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    add_paragraph(doc, "素质目标:", bold=True)
    p = doc.add_paragraph()
    p.add_run("{%p for q in lesson.objectives.quality %}")
    add_paragraph(doc, "{{ loop.index }}. {{ q }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")
    doc.add_paragraph()

    add_paragraph(doc, "【教学重难点】", bold=True, indent_first=False)
    add_paragraph(doc, "教学重点:{{ lesson.key_points }}")
    add_paragraph(doc, "教学难点:{{ lesson.difficult_points }}")
    add_paragraph(doc, "解决措施:{{ lesson.key_difficult_solutions }}")
    doc.add_paragraph()

    add_paragraph(doc, "【教学方法】", bold=True, indent_first=False)
    add_paragraph(doc, "{{ lesson.teaching_methods }}")
    doc.add_paragraph()

    add_paragraph(doc, "【信息化平台及资源】", bold=True, indent_first=False)
    add_paragraph(doc, "{{ lesson.information_resources }}")
    doc.add_paragraph()

    # 三、教学实施过程
    add_section_title(doc, "三、教学实施过程", level=3)

    # 课前 - 4-row pattern: header + {%tr for%} + data + {%tr endfor%}
    add_paragraph(doc, "教学过程:课前", bold=True, indent_first=False)
    _add_steps_table(doc, "lesson.pre_class_steps", include_political=False)
    doc.add_paragraph()

    # 课中 (含思政点)
    add_paragraph(doc, "教学过程:课中", bold=True, indent_first=False)
    _add_steps_table(doc, "lesson.in_class_steps", include_political=True)
    doc.add_paragraph()

    # 课后
    add_paragraph(doc, "教学过程:课后", bold=True, indent_first=False)
    _add_steps_table(doc, "lesson.after_class_steps", include_political=False)
    doc.add_paragraph()

    # 四、教学反思
    add_section_title(doc, "四、教学反思与改进", level=3)
    add_paragraph(doc, "{{ lesson.reflection }}")
    doc.add_paragraph()

    # End loop and page break between lessons
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    output = OUTPUT_DIR / "参赛教案_模板.docx"
    doc.save(str(output))
    print(f"[OK] Saved: {output}")


# ============================================================================
# Template 2: 教学实施报告_模板.docx
# ============================================================================

def build_report_template() -> None:
    doc = Document()

    # ----- Cover -----
    for _ in range(3):
        doc.add_paragraph()
    add_centered_heading(doc, "{{ competition_year }}{{ competition_region }}职业院校技能大赛", font_size=20)
    add_centered_heading(doc, "教学能力比赛", font_size=22)
    for _ in range(2):
        doc.add_paragraph()
    add_centered_heading(doc, "教 学 实 施 报 告", font_size=36)
    for _ in range(3):
        doc.add_paragraph()

    add_label_value_table(doc, [
        ("作品名称：", "{{ work_name }}"),
        ("课程名称：", "{{ course_name }}"),
        ("专业大类：", "{{ major_category }}"),
        ("参赛组别：", "{{ group_name }}"),
    ])
    doc.add_page_break()

    # ----- Title + Intro -----
    add_centered_heading(doc, "{{ course_name }}", font_size=20)
    add_centered_heading(doc, "教学实施报告", font_size=22)
    doc.add_paragraph()
    add_paragraph(doc, "{{ intro_summary }}")
    doc.add_paragraph()

    # ----- 一、整体教学设计 -----
    add_section_title(doc, "一、整体教学设计", level=1)
    doc.add_paragraph()

    add_section_title(doc, "(一) 教学内容分析", level=2)
    add_paragraph(doc, "{{ content_analysis }}")
    add_paragraph(doc, "【图 1:教学内容模块图】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(二) 学情分析", level=2)
    add_paragraph(doc, "{{ student_analysis }}")
    add_paragraph(doc, "【图 2:学情分析模块图】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(三) 教学目标与重难点", level=2)
    add_paragraph(doc, "{{ teaching_objectives }}")
    add_paragraph(doc, "【图 3:教学目标与重难点模块图】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(四) 教学策略", level=2)
    add_paragraph(doc, "{{ teaching_strategy }}")
    add_paragraph(doc, "【图 4:教学策略模块图】", indent_first=False)
    doc.add_paragraph()

    # ----- 二、教学实施过程 -----
    add_section_title(doc, "二、教学实施过程", level=1)
    doc.add_paragraph()

    add_section_title(doc, "(一) 实施过程", level=2)
    impl_table = doc.add_table(rows=4, cols=3)
    impl_table.style = "Table Grid"
    headers = impl_table.rows[0].cells
    headers[0].text = "阶段"
    headers[1].text = "任务"
    headers[2].text = "效果"
    for c in headers:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_background(c, "DEEAF6")
    # 4-row pattern: header / {%tr for%} / data / {%tr endfor%}
    impl_table.rows[1].cells[0].text = "{%tr for stage in implementation_stages %}"
    impl_table.rows[2].cells[0].text = "{{ stage.stage }}"
    impl_table.rows[2].cells[1].text = "{{ stage.task }}"
    impl_table.rows[2].cells[2].text = "{{ stage.effect }}"
    impl_table.rows[3].cells[0].text = "{%tr endfor %}"
    doc.add_paragraph()

    add_section_title(doc, "(二) 混合式教学实施改进", level=2)
    add_paragraph(doc, "{{ blended_teaching_improvement }}")
    add_paragraph(doc, "【图 5:混合式教学模块图】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(三) 多元的评价方式", level=2)

    add_paragraph(doc, "1. 学生评价", bold=True, indent_first=False)
    add_paragraph(doc, "{{ evaluation.student_evaluation }}")
    doc.add_paragraph()

    add_paragraph(doc, "2. 教师评价", bold=True, indent_first=False)
    add_paragraph(doc, "{{ evaluation.teacher_evaluation }}")
    doc.add_paragraph()

    add_paragraph(doc, "3. 企业评价", bold=True, indent_first=False)
    add_paragraph(doc, "{{ evaluation.enterprise_evaluation }}")
    doc.add_paragraph()

    # ----- 三、学生学习效果 -----
    add_section_title(doc, "三、学生学习效果", level=1)
    doc.add_paragraph()

    add_section_title(doc, "(一) 课前学习和拓展学习数量提高,提升学习兴趣", level=2)
    add_paragraph(doc, "{{ learning_effect.pre_class_improvement }}")
    add_paragraph(doc, "【图 6:课前学习数据图】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(二) 课中教学和项目实践质量提高,达成教学目标", level=2)
    add_paragraph(doc, "{{ learning_effect.in_class_improvement }}")
    add_paragraph(doc, "【图 7:学生实践完成情况】", indent_first=False)
    doc.add_paragraph()

    add_section_title(doc, "(三) 取证比例和职业竞赛成绩提高,加强实践能力", level=2)
    add_paragraph(doc, "{{ learning_effect.certification_competition }}")
    add_paragraph(doc, "【图 8:取证及竞赛成绩图】", indent_first=False)
    doc.add_paragraph()

    # ----- 四、教学反思与改进 -----
    add_section_title(doc, "四、教学反思与改进", level=1)
    doc.add_paragraph()

    add_section_title(doc, "(一) 特色创新", level=2)
    p = doc.add_paragraph()
    p.add_run("{%p for f in reflection.feature_innovations %}")
    add_paragraph(doc, "特色 {{ loop.index }}:{{ f }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")
    doc.add_paragraph()

    add_section_title(doc, "(二) 不足", level=2)
    p = doc.add_paragraph()
    p.add_run("{%p for s in reflection.shortcomings %}")
    add_paragraph(doc, "反思 {{ loop.index }}:{{ s }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")
    doc.add_paragraph()

    add_section_title(doc, "(三) 改进措施", level=2)
    p = doc.add_paragraph()
    p.add_run("{%p for m in reflection.improvement_measures %}")
    add_paragraph(doc, "措施 {{ loop.index }}:{{ m }}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    output = OUTPUT_DIR / "教学实施报告_模板.docx"
    doc.save(str(output))
    print(f"[OK] Saved: {output}")


if __name__ == "__main__":
    build_lesson_plan_template()
    build_report_template()
    print("\nAll competition templates generated successfully.")
