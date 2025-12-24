"""Test the fixed document renderer."""
import sys
sys.path.insert(0, 'backend')

from services.document_renderer import DocumentRenderer
from pathlib import Path

# Sample data matching the template structure
test_data = {
    "teaching_topic": "信息安全基础",
    "class_name": "大二",
    "subject": "信息安全技术",
    "grade": "大二",
    "topic": "信息安全基础",
    "duration": "2课时",
    "week_number": "第5周",
    "location": "实训室A205",
    "references": "《信息安全原理》第三版",
    "ideological_political": "培养学生的网络安全意识和社会责任感",
    
    # Teaching goals with structure expected by template
    "teaching_goals": {
        "knowledge": [
            "理解信息安全的基本概念（机密性、完整性、可用性）",
            "掌握常见的信息安全威胁类型",
            "了解信息安全防护的基本策略"
        ],
        "ability": [
            "能够识别常见的安全漏洞",
            "能够应用基本的安全防护措施",
            "能够分析简单的安全事件"
        ],
        "emotion": [
            "培养良好的网络安全意识",
            "树立信息安全责任感",
            "增强网络安全法律法规意识"
        ]
    },
    
    "key_points": "信息安全三要素（CIA）的理解与应用；常见安全威胁的识别",
    "difficult_points": "信息安全防护策略的综合应用；安全事件的分析与处理",
    
    # Teaching steps 
    "teaching_steps": [
        {
            "stage": "课前准备",
            "duration": "5分钟",
            "teacher_activity": "检查教学设备，准备课件和演示环境",
            "student_activity": "签到，准备学习用品",
            "design_intent": "确保教学顺利进行"
        },
        {
            "stage": "导入新课",
            "duration": "10分钟",
            "teacher_activity": "通过真实案例引出信息安全话题",
            "student_activity": "观看案例视频，思考信息安全的重要性",
            "design_intent": "激发学生学习兴趣"
        },
        {
            "stage": "知识讲解",
            "duration": "30分钟",
            "teacher_activity": "讲解CIA三要素，分析常见威胁",
            "student_activity": "记录笔记，参与讨论",
            "design_intent": "系统学习信息安全基础知识"
        }
    ],
    
    "homework": {
        "required": "完成信息安全案例分析报告",
        "optional": "阅读《网络安全法》相关章节"
    },
    
    "reflection": "本次课程学生参与度高，对实际案例很感兴趣",
}

print("🧪 Testing Document Renderer with docxtpl\n")
print("="*80)

# Initialize renderer
renderer = DocumentRenderer()

# Use the template
template_path = "storage/templates/yunnan_forestry_college_template.docx"

print(f"\n📄 Template: {template_path}")
print(f"✅ Template exists: {Path(template_path).exists()}\n")

try:
    print("🔄 Rendering document...")
    output_path = renderer.render_new_document(
        template_path=template_path,
        data=test_data,
        output_filename="test_render_fixed.docx"
    )
    
    print(f"✅ SUCCESS! Document rendered to: {output_path}\n")
    
    # Verify structure
    from docx import Document
    output_doc = Document(output_path)
    template_doc = Document(template_path)
    
    print("📊 Structure Comparison:")
    print(f"   Paragraphs: Template={len(template_doc.paragraphs)}, Output={len(output_doc.paragraphs)}")
    print(f"   Tables: Template={len(template_doc.tables)}, Output={len(output_doc.tables)}")
    
    for i in range(min(len(template_doc.tables), len(output_doc.tables))):
        t_table = template_doc.tables[i]
        o_table = output_doc.tables[i]
        t_size = f"{len(t_table.rows)}x{len(t_table.columns)}"
        o_size = f"{len(o_table.rows)}x{len(o_table.columns)}"
        match = "✅" if t_size == o_size else "❌"
        print(f"   Table {i}: Template={t_size}, Output={o_size} {match}")
    
    print("\n" + "="*80)
    print("🎉 Test completed successfully!")
    
except Exception as e:
    print(f"❌ ERROR: {e}")
    import traceback
    traceback.print_exc()
