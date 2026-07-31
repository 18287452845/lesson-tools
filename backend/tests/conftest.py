"""
Pytest configuration and fixtures for backend testing.

This module provides shared fixtures and configuration for all backend tests.
"""
import pytest
import asyncio
import tempfile
import shutil
from pathlib import Path
from typing import AsyncGenerator, Generator
import aiosqlite
import sys
import os

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from backend.models.database import Database
from backend.config import settings


# ==========================================
# Test Configuration
# ==========================================

@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="session")
def test_data_dir():
    """Get the path to the test data directory."""
    return Path(__file__).parent / "test_data"


# ==========================================
# Temporary Directories
# ==========================================

@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for test files."""
    temp_path = Path(tempfile.mkdtemp())
    yield temp_path
    # Cleanup
    if temp_path.exists():
        shutil.rmtree(temp_path)


@pytest.fixture
def temp_db_path(temp_dir: Path) -> Path:
    """Create a temporary database file path."""
    return temp_dir / "test.db"


# ==========================================
# Database Fixtures
# ==========================================

@pytest.fixture
async def test_db(temp_db_path: Path) -> AsyncGenerator[Database, None]:
    """Create a test database instance with all tables initialized."""
    db = Database(str(temp_db_path))
    await db.initialize()
    yield db
    # Cleanup not needed as temp_dir handles it


@pytest.fixture
async def db_connection(test_db: Database) -> AsyncGenerator[aiosqlite.Connection, None]:
    """Get a database connection for testing."""
    async with test_db.get_connection() as conn:
        yield conn


@pytest.fixture
async def clean_db(test_db: Database) -> Database:
    """Get a database with all tables cleared."""
    async with test_db.get_connection() as conn:
        # Clear all tables
        tables = [
            "lesson_plans", "batch_lesson_plans", "batch_tasks",
            "course_chapter_templates", "template_versions",
            "templates", "classes", "edit_logs", "document_edits", "user_settings"
        ]
        for table in tables:
            await conn.execute(f"DELETE FROM {table}")
        await conn.commit()
    return test_db


# ==========================================
# Test Data Fixtures
# ==========================================

@pytest.fixture
def sample_template_data():
    """Sample template data for testing."""
    return {
        "id": "test-template-1",
        "name": "Test Template",
        "description": "A template for testing",
        "subject": "数学",
        "grade": "一年级",
        "file_path": "/path/to/template.docx",
        "fields_config": '{"fields": ["subject", "grade", "topic"]}',
        "use_count": 0
    }


@pytest.fixture
def sample_lesson_plan_data():
    """Sample lesson plan data for testing."""
    return {
        "id": "test-lesson-1",
        "template_id": "test-template-1",
        "title": "Test Lesson Plan",
        "subject": "数学",
        "grade": "一年级",
        "topic": "加法入门",
        "input_data": '{"key": "value"}',
        "generated_content": '{"content": "test content"}',
        "status": "draft"
    }


@pytest.fixture
def sample_batch_task_data():
    """Sample batch task data for testing."""
    return {
        "id": "test-batch-1",
        "course_name": "数学课程",
        "grade_level": "一年级",
        "total_hours": 48,
        "hours_per_lesson": 2,
        "status": "pending"
    }


@pytest.fixture
def sample_class_data():
    """Sample class data for testing."""
    return {
        "id": "test-class-1",
        "name": "一年级1班",
        "subject": "数学"
    }


# ==========================================
# AI Provider Fixtures
# ==========================================

@pytest.fixture
def mock_ai_response():
    """Mock AI response for testing."""
    return {
        "teaching_goals": ["掌握加法基本概念", "能够进行简单加法运算"],
        "key_points": ["加法的定义", "加法符号的认识"],
        "difficult_points": ["进位概念的理解"],
        "teaching_steps": [
            {"step": 1, "content": "导入新课", "duration": "5分钟"},
            {"step": 2, "content": "讲解新知识", "duration": "15分钟"}
        ],
        "teaching_tools": ["数字卡片", "教学挂图"],
        "homework": "完成课后练习第1-5题",
        "blackboard_design": "板书设计内容",
        "reflection": "教学反思内容"
    }


# ==========================================
# File Fixtures
# ==========================================

@pytest.fixture
def sample_docx_template(temp_dir: Path) -> Path:
    """Create a sample DOCX template file for testing."""
    from docx import Document

    template_path = temp_dir / "test_template.docx"
    doc = Document()

    # Add title
    doc.add_heading("教案模板", 0)

    # Add table with Jinja2 variables
    table = doc.add_table(rows=5, cols=2)
    table.rows[0].cells[0].text = "科目"
    table.rows[0].cells[1].text = "{{ subject }}"
    table.rows[1].cells[0].text = "年级"
    table.rows[1].cells[1].text = "{{ grade }}"
    table.rows[2].cells[0].text = "课题"
    table.rows[2].cells[1].text = "{{ topic }}"
    table.rows[3].cells[0].text = "课时"
    table.rows[3].cells[1].text = "{{ duration }}"

    # Add teaching goals section with loop
    doc.add_heading("教学目标", 1)
    doc.add_paragraph("{% for goal in teaching_goals %}")
    doc.add_paragraph("{{ goal }}")
    doc.add_paragraph("{% endfor %}")

    doc.save(str(template_path))
    return template_path


# ==========================================
# API Client Fixtures
# ==========================================

@pytest.fixture
async def test_client():
    """Create a test HTTP client for API testing."""
    from httpx import AsyncClient, ASGITransport
    from backend.main import app

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        yield client


@pytest.fixture
async def authenticated_client(test_client):
    """Create an authenticated test client (when auth is implemented)."""
    # For now, just return the regular client
    # TODO: Add authentication when implemented
    yield test_client


# ==========================================
# Environment Setup
# ==========================================

@pytest.fixture(scope="session", autouse=True)
def setup_test_environment():
    """Set up test environment variables before all tests."""
    # Set test environment
    os.environ.setdefault("AI_PROVIDER", "deepseek")
    os.environ.setdefault("AI_MAX_TOKENS", "1000")
    os.environ.setdefault("DATABASE_PATH", ":memory:")

    yield

    # Cleanup after all tests
    pass


# ==========================================
# Mock Fixtures
# ==========================================

@pytest.fixture
def mock_ai_generation(monkeypatch):
    """Mock AI generation to avoid API calls during tests."""
    async def mock_generate(*args, **kwargs):
        from backend.models.schemas import GeneratedContent, TeachingGoal, TeachingStep, Homework
        return GeneratedContent(**{
            "teaching_goals": TeachingGoal(knowledge=["目标1", "目标2"]),
            "key_points": "重点1、重点2",
            "difficult_points": "难点1",
            "teaching_steps": [
                TeachingStep(
                    stage="导入",
                    duration="5分钟",
                    teacher_activity="活动内容",
                    student_activity="学生活动",
                    design_intent="设计意图"
                )
            ],
            "teaching_tools": "教具1、教具2",
            "teaching_methods": "方法1、方法2",
            "student_analysis": "学情分析",
            "textbook_analysis": "教材分析",
            "homework": Homework(required="作业内容"),
            "blackboard_design": "板书设计",
            "reflection": "教学反思"
        })

    # Apply mock
    import backend.services.ai_generator
    monkeypatch.setattr(
        backend.services.ai_generator.AIGenerator,
        "generate_lesson_plan",
        mock_generate
    )


# ==========================================
# Skip Markers
# ==========================================

def pytest_configure(config):
    """Configure pytest with custom markers."""
    config.addinivalue_line(
        "markers", "slow: marks tests as slow (deselect with '-m \"not slow\"')"
    )
    config.addinivalue_line(
        "markers", "ai: marks tests that call AI providers (require API keys)"
    )
    config.addinivalue_line(
        "markers", "integration: marks tests as integration tests"
    )
