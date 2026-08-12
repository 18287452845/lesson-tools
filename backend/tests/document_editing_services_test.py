from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from backend.services.document_modifier import (
    DocumentModifier,
    ParagraphFormat,
    add_section_to_document,
    modify_document,
)
from backend.services.lesson_plan_parser import (
    LessonPlanParser,
    SectionInfo,
    get_document_summary,
    parse_document,
)

def _make_lesson_doc(path: Path) -> Path:
    doc = Document()
    goal = doc.add_paragraph()
    run = goal.add_run("教学目标：掌握变量")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True
    goal.paragraph_format.space_before = Pt(3)
    doc.add_paragraph("能够解决实际问题")
    doc.add_paragraph("")
    doc.add_paragraph("教学过程：")
    doc.add_paragraph("导入新课：问题情境")
    doc.add_paragraph("学生观察")
    doc.add_paragraph("探究新知：变量概念")
    doc.add_paragraph("教师讲解")
    doc.add_paragraph("课堂总结：归纳")
    doc.add_paragraph("教学反思：课后填写")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "教学重点"
    table.cell(0, 1).text = "变量与数据类型"
    table.cell(1, 0).text = "教学难点"
    table.cell(1, 1).text = "类型转换"

    single = doc.add_table(rows=1, cols=1)
    single.cell(0, 0).text = "教学方法：合作探究"
    doc.save(path)
    return path


def test_parser_discovers_paragraph_table_and_missing_sections(tmp_path):
    path = _make_lesson_doc(tmp_path / "lesson.docx")
    parser = LessonPlanParser(path)

    sections = parser.parse()

    assert sections["teaching_goals"].content == "掌握变量\n能够解决实际问题"
    assert sections["teaching_goals"].start_para_idx == 0
    assert sections["key_points"].in_table is True
    assert sections["key_points"].content == "变量与数据类型"
    assert sections["teaching_methods"].content == "合作探究"
    assert sections["homework"].found is False
    assert parser.get_section_content("key_points") == "变量与数据类型"
    assert parser.get_section_content("missing") is None
    assert parser.is_section_found("teaching_process") is True
    assert parser.is_section_found("missing") is False

    parsed = parser.get_parsed_sections()
    assert parsed["key_points"].location.table_idx == 0
    assert parsed["homework"].location is None

    summary = parser.get_document_summary()
    assert summary["filename"] == "lesson.docx"
    assert summary["paragraph_count"] == 10
    assert summary["table_count"] == 2
    assert summary["sections_found"] >= 6

    stages = parser.parse_teaching_process_stages()
    assert [stage["stage"] for stage in stages] == ["问题情境", "变量概念", "归纳"]
    assert stages[0]["content"] == ["学生观察"]
    assert stages[1]["content"] == ["教师讲解"]


def test_parser_helpers_and_convenience_functions(tmp_path):
    path = _make_lesson_doc(tmp_path / "helpers.docx")
    parser = LessonPlanParser(path)
    assert parser.get_document_summary() == {}
    assert parser.parse_teaching_process_stages() == []
    parser._parse_paragraphs()
    assert parser.sections == {}
    assert parser._match_section_header("Learning Objectives: fractions") == "teaching_goals"
    assert parser._match_section_header("ordinary text") is None
    assert parser._extract_content_after_header("重点：理解概念", "key_points") == "理解概念"
    assert parser._extract_content_after_header("重点：", "key_points") is None
    assert parser._extract_content_after_header("anything", "unknown") is None

    parser.doc = Document(str(path))
    table = parser.doc.tables[0]
    assert parser._extract_adjacent_content(table, 0, 0) == "变量与数据类型"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "这是一段足够长的下方内容，用于验证解析器会发现跨行布局中的正文"
    assert "下方内容" in parser._extract_adjacent_content(table, 0, 0)
    assert parser._extract_adjacent_content(table, 1, 1) is None

    parser.sections["teaching_process"] = SectionInfo(
        name="teaching_process",
        found=True,
        content="自定义环节：准备\n继续活动",
    )
    assert parser.parse_teaching_process_stages() == [
        {"stage": "自定义环节", "content": ["准备", "继续活动"]}
    ]

    assert parse_document(str(path))["teaching_goals"].found is True
    assert get_document_summary(str(path))["table_count"] == 2
    with pytest.raises(FileNotFoundError):
        LessonPlanParser(tmp_path / "absent.docx").parse()


def test_modifier_replaces_paragraphs_and_table_cells(tmp_path):
    path = _make_lesson_doc(tmp_path / "edit.docx")
    parser = LessonPlanParser(path)
    sections = parser.parse()
    modifier = DocumentModifier(str(path))

    assert modifier.modify_section(sections["teaching_goals"], "新目标\n新续行") is True
    assert modifier.doc.paragraphs[0].text == "教学目标：新目标"
    assert modifier.doc.paragraphs[1].text == "新续行"
    assert modifier.doc.paragraphs[0].runs[0].font.size.pt == 11
    assert modifier.modify_section(sections["key_points"], "更新后的重点") is True
    assert modifier.doc.tables[0].cell(0, 1).text == "更新后的重点"
    assert modifier.modify_section(SectionInfo("none", False, None), "x") is False

    no_indexes = SectionInfo("bad", True, "old")
    assert modifier._modify_paragraph_section(no_indexes, "x", True) is False
    assert modifier._modify_table_section(no_indexes, "x", True) is False
    bad_table = SectionInfo("bad", True, "old", in_table=True, table_idx=99, cell_location=(0, 0))
    assert modifier._modify_table_section(bad_table, "x", True) is False
    bad_cell = SectionInfo("bad", True, "old", in_table=True, table_idx=0, cell_location=(99, 0))
    assert modifier._modify_table_section(bad_cell, "x", True) is False

    plain = SectionInfo("plain", True, "普通内容", start_para_idx=3, end_para_idx=3)
    assert modifier._modify_paragraph_section(plain, "无标题替换", False) is True
    assert modifier.doc.paragraphs[3].text == "无标题替换"
    assert len(modifier.edit_history) == 2


def test_modifier_append_insert_add_save_undo_and_wrappers(tmp_path):
    path = _make_lesson_doc(tmp_path / "operations.docx")
    parser = LessonPlanParser(path)
    goals = parser.parse()["teaching_goals"]

    modifier = DocumentModifier(str(path))
    assert modifier.undo() is False
    empty = SectionInfo("empty", True, "", start_para_idx=0, end_para_idx=0)
    assert modifier.append_to_section(empty, "x") is False
    assert modifier.insert_in_section(empty, "x") is False
    assert modifier.append_to_section(goals, "追加内容", preserve_format=False) is True
    assert "追加内容" in "\n".join(p.text for p in modifier.doc.paragraphs)

    modifier = DocumentModifier(str(path))
    assert modifier.insert_in_section(goals, "开头", at_start=True, preserve_format=False) is True
    assert "开头" in modifier.doc.paragraphs[0].text
    modifier = DocumentModifier(str(path))
    assert modifier.insert_in_section(goals, "结尾", at_start=False, preserve_format=False) is True
    assert "结尾" in "\n".join(p.text for p in modifier.doc.paragraphs)

    style = ParagraphFormat(
        font_name="Arial",
        font_size=12,
        bold=True,
        italic=True,
        underline=True,
        color=RGBColor(1, 2, 3),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_before=2,
        space_after=3,
    )
    assert modifier.add_section("homework", "完成练习", format_style=style) is True
    assert modifier.add_section("custom", "扩展", "after_section", "homework") is True
    assert modifier._get_section_header("custom") == "custom："
    assert modifier.doc.paragraphs[-4].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert modifier.undo() is True

    explicit = tmp_path / "saved.docx"
    assert modifier.save(str(explicit)) == str(explicit)
    assert explicit.exists()
    default = modifier.save()
    assert default.endswith("operations_edited.docx")
    suffixed = modifier.save_with_suffix("_copy")
    assert suffixed.endswith("operations_copy.docx")

    modified = modify_document(str(path), goals, "包装器目标")
    assert Path(modified).exists()
    added = add_section_to_document(str(path), "homework", "包装器作业")
    assert Path(added).exists()
    with pytest.raises(FileNotFoundError):
        DocumentModifier(str(tmp_path / "missing.docx"))
