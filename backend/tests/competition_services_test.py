import json
from pathlib import Path

import pytest
from docx import Document

from backend.models.schemas import (
    CompetitionLessonContent,
    CompetitionOverallDesign,
    CompetitionProject,
    CompetitionReportContent,
    CompetitionSingleLesson,
)
from backend.services import competition_generator as generator_module
from backend.services import competition_renderer as renderer_module
from backend.services.competition_generator import CompetitionGenerator, _project_context
from backend.services.competition_renderer import CompetitionRenderer


def _project(**updates):
    data = {
        "id": "project-1",
        "name": "参赛项目",
        "competition_year": "2026年",
        "competition_region": "云南省",
        "competition_level": "省赛",
        "work_name": "Python/安全 实训",
        "course_name": "Python程序设计",
        "major_category": "电子信息",
        "major_name": "计算机应用",
        "group_name": "专业组",
        "total_hours": 4,
        "hours_per_lesson": 2,
        "class_name": "计应1班",
        "location": "实训室",
        "textbook_info": {"name": "Python教材"},
        "context_data": {"level": "初级"},
        "created_at": "2026-01-01T00:00:00",
        "updated_at": "2026-01-01T00:00:00",
    }
    data.update(updates)
    return CompetitionProject(**data)


def _lesson_content():
    return CompetitionLessonContent(
        overall_design=CompetitionOverallDesign(
            content_analysis="**内容**分析",
            teaching_method="任务驱动",
            survey_questions=["问题1"],
        ),
        lessons=[
            CompetitionSingleLesson(lesson_number=1, title="环境搭建"),
            CompetitionSingleLesson(lesson_number=2, title="变量应用"),
        ],
    )


def _report_content():
    return CompetitionReportContent(
        intro_summary="实施总述",
        implementation_stages=[{"stage": "课前", "task": "预习", "effect": "完成"}],
        reflection={"feature_innovations": ["创新"], "shortcomings": [], "improvement_measures": []},
    )


@pytest.mark.asyncio
async def test_competition_generator_json_retry_topics_lesson_and_report(monkeypatch):
    project = _project()
    context = _project_context(project)
    assert "Python程序设计" in context and "Python教材" in context and "初级" in context

    generator = CompetitionGenerator(max_concurrent_lessons=1)
    assert generator._strip_code_fences("```json\n{\"ok\": true}\n```") == '{"ok": true}'
    assert generator._parse_json("说明 {\"ok\": true} 尾声") == {"ok": True}
    with pytest.raises(ValueError, match="did not contain"):
        generator._parse_json("not json")
    with pytest.raises(ValueError, match="Failed to parse"):
        generator._parse_json("prefix {bad} suffix")

    calls = []

    async def fake_generate(**kwargs):
        calls.append(kwargs["prompt"])
        return "invalid" if len(calls) == 1 else '{"ok": 1}'

    async def no_sleep(_):
        return None

    monkeypatch.setattr(generator_module, "generate_with_ai", fake_generate)
    monkeypatch.setattr(generator_module.asyncio, "sleep", no_sleep)
    assert await generator._ai_call_json("prompt", "system", max_attempts=2) == {"ok": 1}
    assert "仅输出严格的 JSON" in calls[1]
    with pytest.raises(ValueError):
        await generator._ai_call_json("prompt", "system", max_attempts=0)

    async def fake_json(prompt, system_prompt, max_attempts=3):
        if "数组字段 topics" in prompt:
            return {"topics": ["AI任务一"]}
        if "【整体设计】" in prompt:
            return {
                "content_analysis": "内容",
                "student_analysis": "学情",
                "goal_analysis": "目标",
                "process_design": "过程",
                "teaching_method": "方法",
                "survey_questions": ["问题"],
            }
        if "完整内容" in prompt:
            return {
                "module_name": "模块",
                "student_analysis": {"knowledge_basis": "基础"},
                "objectives": {"knowledge": ["知识"], "ability": ["能力"], "quality": ["素质"]},
                "pre_class_steps": [{"stage": "预习"}],
                "in_class_steps": [{"stage": "实践", "duration": "80min"}],
                "after_class_steps": [{"stage": "拓展"}],
            }
        if "整体教学设计" in prompt:
            return {
                "intro_summary": "总述",
                "content_analysis": "内容",
                "implementation_stages": [{"stage": "课前", "task": "预习", "effect": "良好"}],
                "evaluation": {"student_evaluation": "学生评价"},
            }
        return {
            "learning_effect": {"pre_class_improvement": "提升"},
            "reflection": {
                "feature_innovations": ["创新"],
                "shortcomings": ["不足"],
                "improvement_measures": ["改进"],
            },
        }

    monkeypatch.setattr(generator, "_ai_call_json", fake_json)
    progress = []

    async def progress_cb(current, total, message):
        progress.append((current, total, message))

    content = await generator.generate_full_lesson_plan(
        project,
        topics_input="- 用户任务",
        additional_requirements="体现课程思政",
        progress_callback=progress_cb,
    )
    assert [lesson.title for lesson in content.lessons] == ["AI任务一", "任务2"]
    assert content.lessons[0].project_name == project.work_name
    assert content.lessons[0].objectives.knowledge == ["知识"]
    assert progress[-1][:2] == (4, 4)

    assert await generator._resolve_topics(project, "任务一\n任务二\n任务三", 2) == ["任务一", "任务二"]
    with pytest.raises(ValueError, match="invalid topics"):
        old = generator._ai_call_json

        async def invalid_topics(*args, **kwargs):
            return {"topics": "bad"}

        monkeypatch.setattr(generator, "_ai_call_json", invalid_topics)
        await generator._resolve_topics(project, None, 2)
    monkeypatch.setattr(generator, "_ai_call_json", old)

    report = await generator.generate_report(
        project,
        related_lesson_plan=content,
        additional_requirements="数据量化",
        progress_callback=progress_cb,
    )
    assert report.implementation_stages[0].task == "预习"
    assert report.learning_effect.pre_class_improvement == "提升"
    assert "AI任务一" in generator._build_related_context(content)
    assert generator._build_related_context(None) == ""

    async def broken_progress(*args):
        raise RuntimeError("callback failed")

    await generator.generate_report(project, progress_callback=broken_progress)


def _write_template(path: Path, text: str):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_competition_renderer_lesson_report_helpers_and_missing_templates(tmp_path, monkeypatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    main = tmp_path / "main.docx"
    single = tmp_path / "single.docx"
    report_template = tmp_path / "report.docx"
    _write_template(main, "{{ work_name }} {{ overall_design.content_analysis }}")
    _write_template(single, "{{ lesson.lesson_number }} {{ lesson.title }}")
    _write_template(report_template, "{{ work_name }} {{ intro_summary }}")

    monkeypatch.setattr(renderer_module, "LESSON_PLAN_MAIN_TEMPLATE", main)
    monkeypatch.setattr(renderer_module, "LESSON_PLAN_SINGLE_TEMPLATE", single)
    monkeypatch.setattr(renderer_module, "REPORT_TEMPLATE", report_template)

    renderer = CompetitionRenderer()
    renderer.output_dir = output_dir
    project = _project()
    lesson_path = Path(renderer.render_lesson_plan(project, _lesson_content()))
    report_path = Path(renderer.render_report(project, _report_content()))
    assert lesson_path.exists() and report_path.exists()
    assert len(Document(lesson_path).paragraphs) >= 5
    assert "实施总述" in "\n".join(p.text for p in Document(report_path).paragraphs)
    assert not list(output_dir.glob("_temp_*.docx"))

    name = renderer._build_output_filename("参赛教案", "非法 / 名称", "now")
    assert name == "参赛教案_非法_名称_now.docx"
    context = renderer._project_to_context(project)
    assert context["total_hours"] == 4 and context["work_name"] == project.work_name
    assert renderer._clean_dict({"a": None, "b": ["**文本**", 2]}) == {"a": "", "b": ["文本", 2]}

    missing = tmp_path / "missing.docx"
    monkeypatch.setattr(renderer_module, "LESSON_PLAN_MAIN_TEMPLATE", missing)
    with pytest.raises(FileNotFoundError, match="Main template"):
        renderer.render_lesson_plan(project, _lesson_content())
    monkeypatch.setattr(renderer_module, "LESSON_PLAN_MAIN_TEMPLATE", main)
    monkeypatch.setattr(renderer_module, "LESSON_PLAN_SINGLE_TEMPLATE", missing)
    with pytest.raises(FileNotFoundError, match="Single-lesson"):
        renderer.render_lesson_plan(project, _lesson_content())
    monkeypatch.setattr(renderer_module, "REPORT_TEMPLATE", missing)
    with pytest.raises(FileNotFoundError, match="Report template"):
        renderer.render_report(project, _report_content())
