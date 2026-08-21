import zipfile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from backend.services import document_renderer


def test_clean_text_removes_blank_lines_without_inventing_line_breaks():
    renderer = document_renderer.DocumentRenderer()

    text = "第一项完成。  2. 第二项继续\r\n\r\n  【教师】讲解  \n   \n【学生】练习"

    assert renderer._clean_text_for_output(text) == (
        "第一项完成。 2. 第二项继续\n【教师】讲解\n【学生】练习"
    )


def test_markdown_cleaning_preserves_python_symbols_and_cleans_media():
    renderer = document_renderer.DocumentRenderer()

    text = (
        "公式 s=3.14*r*r，构造函数 __init__，**重点**，_斜体_；"
        "![流程图](https://example.com/flow.png)，[参考资料](https://example.com)"
    )

    cleaned = renderer._clean_text_for_output(text)

    assert "s=3.14*r*r" in cleaned
    assert "__init__" in cleaned
    assert "**重点**" not in cleaned
    assert "_斜体_" not in cleaned
    assert "!流程图" not in cleaned
    assert "流程图" in cleaned
    assert "参考资料" in cleaned
    assert "https://" not in cleaned


def test_render_autoescapes_dynamic_xml_text(tmp_path):
    template_path = tmp_path / "escape-template.docx"
    template = Document()
    template.add_paragraph("比较结果：{{ key_points }}")
    template.save(template_path)

    renderer = document_renderer.DocumentRenderer()
    renderer.output_dir = tmp_path
    output_path = renderer.render_lesson_plan(
        str(template_path),
        {"topic": "转义测试", "key_points": "a < b & c > d"},
    )

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    assert not root.xpath(
        "//w:t/*",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
    )

    rendered = Document(output_path)
    assert "比较结果：a < b & c > d" in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )


def test_combined_render_autoescapes_each_lesson(tmp_path):
    template_path = tmp_path / "combined-escape-template.docx"
    template = Document()
    template.add_paragraph("{{ lesson_title }}：{{ key_points }}")
    template.save(template_path)

    renderer = document_renderer.DocumentRenderer()
    renderer.output_dir = tmp_path
    output_path = renderer.render_lesson_plans_document(
        template_path=str(template_path),
        lesson_plans_data=[
            {"lesson_number": 1, "topic": "比较运算", "key_points": "a < b & c > d"},
            {"lesson_number": 2, "topic": "位运算", "key_points": "x << 1 & y >> 1"},
        ],
        course_name="转义测试",
        document_number=1,
    )

    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    assert not root.xpath(
        "//w:t/*",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
    )

    rendered = Document(output_path)
    rendered_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "教案1：比较运算：a < b & c > d" in rendered_text
    assert "教案2：位运算：x << 1 & y >> 1" in rendered_text


def test_compact_homepage_controls_table_pagination(tmp_path):
    path = tmp_path / "homepage.docx"
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "https://example.com/" + "x" * 200
    process_table = document.add_table(rows=3, cols=2)
    process_table.cell(0, 0).text = "教学过程"
    process_table.cell(0, 1).text = "主要教学内容及步骤"
    process_table.cell(2, 0).text = "作业布置"
    process_table.cell(2, 1).text = "完成课后作业"
    last_row_properties = process_table.rows[-1]._tr.get_or_add_trPr()
    height = etree.SubElement(last_row_properties, qn("w:trHeight"))
    height.set(qn("w:val"), "2809")
    height.set(qn("w:hRule"), "atLeast")
    document.save(path)

    document_renderer.DocumentRenderer._compact_rendered_homepage(str(path))

    rendered = Document(path)
    layout = rendered.tables[0]._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None
    assert layout.get(qn("w:type")) == "fixed"
    for cell in rendered.tables[1].rows[0].cells:
        for paragraph in cell.paragraphs:
            assert paragraph._p.pPr.find(qn("w:keepNext")) is not None
    assert rendered.tables[1].rows[-1]._tr.trPr.find(qn("w:trHeight")) is None
    section_properties = rendered._element.body.sectPr
    terminal_paragraph = section_properties.getprevious()
    assert terminal_paragraph.tag == qn("w:p")
    spacing = terminal_paragraph.pPr.find(qn("w:spacing"))
    assert spacing.get(qn("w:before")) == "0"
    assert spacing.get(qn("w:after")) == "0"
    assert spacing.get(qn("w:line")) == "20"
    assert spacing.get(qn("w:lineRule")) == "exact"
    size = terminal_paragraph.pPr.find("w:rPr/w:sz", terminal_paragraph.nsmap)
    assert size.get(qn("w:val")) == "2"


def test_process_data_adds_role_markers_and_numeric_template_duration():
    renderer = document_renderer.DocumentRenderer()

    processed = renderer._process_data(
        {
            "duration": "2课时",
            "teaching_steps": [
                {
                    "stage": "传授新知",
                    "duration": "30分钟",
                    "teacher_activity": "讲解配置原理",
                    "student_activity": "完成同步操作",
                    "design_intent": "建立知识与操作的联系",
                }
            ],
        }
    )

    assert processed["duration"] == "2"
    assert processed["teaching_steps"][0]["teacher_activity"].startswith("【教师】")
    assert processed["teaching_steps"][0]["student_activity"].startswith("【学生】")


def test_process_data_separates_and_deduplicates_online_resources():
    renderer = document_renderer.DocumentRenderer()
    online = (
        "1. 官方文档：https://example.com/guide.html"
        "2. 代码仓库：https://github.com/example/project"
        "3. 在线课程：https://example.com/course"
    )

    processed = renderer._process_data(
        {
            "references": f"《Python教材》\n{online}",
            "online_resources": online,
        }
    )

    assert processed["references"] == "《Python教材》"
    assert processed["electronic_resources"].splitlines() == [
        "1. 官方文档：https://example.com/guide.html",
        "2. 代码仓库：https://github.com/example/project",
        "3. 在线课程：https://example.com/course",
    ]


def test_process_data_separates_adjacent_urls():
    renderer = document_renderer.DocumentRenderer()

    processed = renderer._process_data(
        {
            "online_resources": (
                "https://example.com/onehttps://example.com/two"
                "https://example.com/three"
            )
        }
    )

    assert processed["electronic_resources"].splitlines() == [
        "https://example.com/one",
        "https://example.com/two",
        "https://example.com/three",
    ]


def test_process_data_does_not_split_digits_inside_resource_url():
    renderer = document_renderer.DocumentRenderer()

    processed = renderer._process_data(
        {
            "online_resources": (
                "1. 中国大学MOOC：https://www.icourse163.org/course/PYTHON-1"
                "2. Python文档：https://docs.python.org/3/"
            )
        }
    )

    assert processed["electronic_resources"].splitlines() == [
        "1. 中国大学MOOC：https://www.icourse163.org/course/PYTHON-1",
        "2. Python文档：https://docs.python.org/3/",
    ]


def test_combine_documents_adds_page_break_and_keeps_sectpr_last(
    temp_dir,
    monkeypatch,
):
    first_path = temp_dir / "first.docx"
    second_path = temp_dir / "second.docx"

    first_doc = Document()
    first_doc.add_paragraph("教案一")
    first_doc.save(first_path)

    second_doc = Document()
    second_doc.add_paragraph("教案二")
    second_doc.save(second_path)

    renderer = document_renderer.DocumentRenderer()
    monkeypatch.setattr(renderer, "output_dir", temp_dir)

    output_path = renderer._combine_documents_with_page_breaks(
        [str(first_path), str(second_path)],
        course_name="测试课程",
        document_number=1,
        week_number=1,
    )

    combined = Document(output_path)
    page_breaks = [
        element
        for element in combined.element.body.iter()
        if element.tag == qn("w:br") and element.get(qn("w:type")) == "page"
    ]

    assert len(page_breaks) == 1
    assert [paragraph.text for paragraph in combined.paragraphs if paragraph.text] == [
        "教案一",
        "教案二",
    ]
    assert combined.element.body[-1].tag == qn("w:sectPr")
