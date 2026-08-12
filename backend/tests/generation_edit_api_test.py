import io
import json
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException, UploadFile

from backend.api import edit as edit_api
from backend.api import generate as generate_api
from backend.models.schemas import (
    AIEnhanceRequest,
    AddSectionRequest,
    FieldEditRequest,
    GeneratedContent,
    LessonPlanGenerateRequest,
    SectionEditRequest,
)


def _request(topic="变量"):
    return LessonPlanGenerateRequest(
        template_id="yunlin-standard",
        subject="Python",
        grade="大学",
        topic=topic,
        duration="2课时",
        textbook_name="Python教材",
        online_resources="在线课程",
        class_ids=["class-1"],
        location="实训室",
        generate_reflection=True,
    )


def _content():
    return GeneratedContent(
        teaching_goals={"knowledge": ["掌握变量"]},
        key_points="变量定义",
        teaching_steps=[
            {
                "stage": "实践",
                "duration": "30分钟",
                "teacher_activity": "指导",
                "student_activity": "操作",
                "design_intent": "应用",
            }
        ],
        homework={"required": "完成练习"},
        online_resources="AI资源",
    )


async def _stream_payloads(response):
    return [json.loads(item.removeprefix("data: ").strip()) async for item in response.body_iterator]


@pytest.fixture
async def api_db(test_db, monkeypatch):
    monkeypatch.setattr(generate_api, "db", test_db)
    monkeypatch.setattr(edit_api, "db", test_db)
    await test_db.execute(
        """INSERT INTO templates
        (id, name, file_path, fields_config, created_at, updated_at)
        VALUES ('yunlin-standard', '云林模板', 'builtin.docx',
                '[{"name":"key_points","display_name":"教学重点","field_type":"text"}]', 'x', 'x')""",
        commit=True,
    )
    await test_db.execute(
        "INSERT INTO classes (id, name, created_at, updated_at) VALUES ('class-1','计应1班','x','x')",
        commit=True,
    )
    return test_db


@pytest.mark.asyncio
async def test_generation_crud_regeneration_update_export_and_stream(api_db, tmp_path, monkeypatch):
    class FakeGenerator:
        async def generate_lesson_plan(self, request, field_configs=None, generate_reflection=False):
            assert generate_reflection is True
            return _content()

        async def regenerate_field(self, input_data, field_name, current_content, instruction):
            if field_name == "teaching_steps":
                return '说明 [{"stage":"新步骤"}] 结束'
            if field_name == "teaching_goals":
                return '```json\n{"knowledge":["新目标"]}\n```'
            return "新文本"

    generator = FakeGenerator()

    async def get_generator():
        return generator

    async def no_sleep(_):
        return None

    monkeypatch.setattr(generate_api, "get_ai_generator", get_generator)
    monkeypatch.setattr(generate_api.asyncio, "sleep", no_sleep)

    created = await generate_api.generate_lesson_plan(_request())
    plan_id = created.id
    assert created.content.key_points == "变量定义"
    assert generate_api.extract_json_from_content('prefix {"a": 1}') == {"a": 1}
    assert generate_api.extract_json_from_content("[1, 2]", True) == [1, 2]

    stream = await generate_api.generate_lesson_plan_stream(_request("流式变量"))
    events = [event async for event in stream.body_iterator]
    assert any('"type": "complete"' in event for event in events)
    assert any('"progress": 60' in event for event in events)

    regenerated = await generate_api.regenerate_field(plan_id, "teaching_steps", "更详细")
    assert regenerated["new_content"] == [{"stage": "新步骤"}]
    regenerated = await generate_api.regenerate_field(plan_id, "teaching_goals")
    assert regenerated["new_content"] == {"knowledge": ["新目标"]}
    regenerated = await generate_api.regenerate_field(plan_id, "key_points")
    assert regenerated["new_content"] == "新文本"

    await generate_api.update_field(
        plan_id,
        FieldEditRequest(lesson_plan_id=plan_id, field_name="key_points", content="手工重点"),
    )
    row = await api_db.fetch_one("SELECT generated_content FROM lesson_plans WHERE id=?", (plan_id,))
    assert json.loads(row["generated_content"])["key_points"] == "手工重点"
    await api_db.execute(
        "UPDATE lesson_plans SET final_content=? WHERE id=?",
        (json.dumps({"key_points": "旧最终内容"}), plan_id),
        commit=True,
    )
    await generate_api.update_field(
        plan_id,
        FieldEditRequest(lesson_plan_id=plan_id, field_name="key_points", content="最终重点"),
    )
    row = await api_db.fetch_one("SELECT final_content FROM lesson_plans WHERE id=?", (plan_id,))
    assert json.loads(row["final_content"])["key_points"] == "最终重点"

    output = tmp_path / "export.docx"

    class FakeRenderer:
        def render_lesson_plan(self, template_path, render_data):
            assert render_data["class_name"] == "计应1班"
            assert render_data["references"] == "《Python教材》\n在线课程"
            output.write_bytes(b"docx")
            return str(output)

    monkeypatch.setattr(generate_api, "DocumentRenderer", FakeRenderer)
    monkeypatch.setattr(generate_api, "get_builtin_template_path", lambda: tmp_path / "template.docx")
    exported = await generate_api.export_lesson_plan(plan_id)
    assert exported.filename == "export.docx"

    details = await generate_api.get_lesson_plan(plan_id)
    assert details["final_content"]["key_points"] == "最终重点"
    listing = await generate_api.list_lesson_plans(
        limit=10, offset=0, subject="Python", grade="大学", status="completed"
    )
    assert listing[0]["id"] == plan_id
    assert len(await generate_api.list_lesson_plans(limit=10, offset=0)) >= 2
    assert (await generate_api.delete_lesson_plan(plan_id))["message"].startswith("Lesson plan deleted")

    for func, args in (
        (generate_api.get_lesson_plan, ("missing",)),
        (generate_api.delete_lesson_plan, ("missing",)),
        (generate_api.regenerate_field, ("missing", "key_points")),
        (
            generate_api.update_field,
            ("missing", FieldEditRequest(lesson_plan_id="missing", field_name="x", content="y")),
        ),
        (generate_api.export_lesson_plan, ("missing",)),
    ):
        with pytest.raises(HTTPException) as exc:
            await func(*args)
        assert exc.value.status_code == 404


@pytest.mark.asyncio
async def test_generation_maps_template_ai_and_stream_errors(api_db, monkeypatch):
    with pytest.raises(HTTPException) as exc:
        await generate_api.generate_lesson_plan(_request().model_copy(update={"template_id": "bad"}))
    assert exc.value.status_code == 400

    await api_db.execute("DELETE FROM templates", commit=True)
    with pytest.raises(HTTPException) as exc:
        await generate_api.generate_lesson_plan(_request())
    assert exc.value.status_code == 404
    missing_stream = await generate_api.generate_lesson_plan_stream(_request())
    assert any(event.get("message") == "模板不存在" for event in await _stream_payloads(missing_stream))

    await api_db.execute(
        "INSERT INTO templates (id,name,file_path,fields_config) VALUES ('yunlin-standard','x','x','bad-json')",
        commit=True,
    )

    class BrokenGenerator:
        error = ValueError("invalid content")

        async def generate_lesson_plan(self, *args, **kwargs):
            raise self.error

        async def regenerate_field(self, *args, **kwargs):
            raise RuntimeError("regenerate failed")

    broken = BrokenGenerator()

    async def get_broken():
        return broken

    monkeypatch.setattr(generate_api, "get_ai_generator", get_broken)
    with pytest.raises(HTTPException) as exc:
        await generate_api.generate_lesson_plan(_request())
    assert exc.value.status_code == 400
    stream = await generate_api.generate_lesson_plan_stream(_request())
    assert any("invalid content" in event.get("message", "") for event in await _stream_payloads(stream))

    broken.error = RuntimeError("offline")
    with pytest.raises(HTTPException) as exc:
        await generate_api.generate_lesson_plan(_request())
    assert exc.value.status_code == 500
    stream = await generate_api.generate_lesson_plan_stream(_request())
    assert any("AI生成失败" in event.get("message", "") for event in await _stream_payloads(stream))

    await api_db.execute(
        """INSERT INTO lesson_plans
        (id,template_id,title,input_data,generated_content,status)
        VALUES ('regen-fail','yunlin-standard','x','{}','{}','generated')""",
        commit=True,
    )
    with pytest.raises(HTTPException) as exc:
        await generate_api.regenerate_field("regen-fail", "key_points")
    assert exc.value.status_code == 500


def _write_edit_doc(path):
    doc = Document()
    doc.add_paragraph("教学目标：掌握变量")
    doc.add_paragraph("教学过程：")
    doc.add_paragraph("导入：案例")
    doc.save(path)


@pytest.mark.asyncio
async def test_edit_upload_operations_ai_add_save_undo_history(api_db, tmp_path, monkeypatch):
    monkeypatch.setattr(edit_api.settings, "upload_dir", tmp_path)
    source = tmp_path / "source.docx"
    _write_edit_doc(source)
    upload = UploadFile(file=io.BytesIO(source.read_bytes()), filename="lesson.docx")
    uploaded = await edit_api.upload_document(upload)
    doc_id = uploaded.id
    assert uploaded.parsed_sections["teaching_goals"].found is True
    assert (await edit_api.get_document(doc_id))["filename"] == "lesson.docx"

    replaced = await edit_api.edit_section(
        doc_id, SectionEditRequest(section_name="teaching_goals", operation="replace", content="新目标")
    )
    assert replaced.new_content == "新目标"
    appended = await edit_api.edit_section(
        doc_id, SectionEditRequest(section_name="teaching_goals", operation="append", content="追加目标")
    )
    assert "追加目标" in appended.new_content
    inserted = await edit_api.edit_section(
        doc_id, SectionEditRequest(section_name="teaching_goals", operation="insert", content="前置目标")
    )
    assert inserted.new_content.startswith("前置目标")

    class FakeEditor:
        async def generate_missing_section(self, section_name, context, instruction=None):
            return "AI生成内容"

        async def modify_content(self, content, instruction, context, section_name):
            return "AI修改内容"

        async def enhance_content(self, content, kind, context, instruction=None):
            return "AI增强内容"

    fake_editor = FakeEditor()

    async def get_editor():
        return fake_editor

    monkeypatch.setattr(edit_api, "get_ai_editor", get_editor)
    monkeypatch.setattr(edit_api, "AIEditor", lambda: fake_editor)
    modified = await edit_api.edit_section(
        doc_id,
        SectionEditRequest(
            section_name="teaching_goals", operation="ai_modify", ai_instruction="更专业"
        ),
    )
    assert modified.new_content == "AI修改内容"
    generated = await edit_api.edit_section(
        doc_id, SectionEditRequest(section_name="teaching_goals", operation="generate")
    )
    assert generated.new_content == "AI生成内容"
    enhanced = await edit_api.ai_enhance_section(
        doc_id, AIEnhanceRequest(section_name="teaching_goals", enhancement_type="professional")
    )
    assert enhanced["enhanced_content"] == "AI增强内容"

    added = await edit_api.add_section(
        doc_id,
        AddSectionRequest(section_name="reflection", ai_generate=False, manual_content="课后反思"),
    )
    assert added["content"] == "课后反思"
    added_ai = await edit_api.add_section(
        doc_id, AddSectionRequest(section_name="homework", ai_generate=True)
    )
    assert added_ai["content"] == "AI生成内容"

    history = await edit_api.get_edit_history(doc_id)
    assert len(history["edits"]) >= 7
    saved = await edit_api.save_document(doc_id)
    assert saved.filename == "lesson_edited.docx"
    assert (await edit_api.undo_edit(doc_id))["success"] is True


@pytest.mark.asyncio
async def test_edit_validation_and_not_found_paths(api_db, tmp_path, monkeypatch):
    monkeypatch.setattr(edit_api.settings, "upload_dir", tmp_path)
    with pytest.raises(HTTPException) as exc:
        await edit_api.upload_document(UploadFile(file=io.BytesIO(b"x"), filename="bad.txt"))
    assert exc.value.status_code == 400

    for func, args in (
        (edit_api.get_document, ("missing",)),
        (edit_api.save_document, ("missing",)),
        (edit_api.undo_edit, ("missing",)),
        (edit_api.get_edit_history, ("missing",)),
        (
            edit_api.edit_section,
            ("missing", SectionEditRequest(section_name="x", content="y")),
        ),
        (
            edit_api.ai_enhance_section,
            ("missing", AIEnhanceRequest(section_name="x")),
        ),
        (
            edit_api.add_section,
            ("missing", AddSectionRequest(section_name="x")),
        ),
    ):
        with pytest.raises(HTTPException) as exc:
            await func(*args)
        assert exc.value.status_code == 404

    source = tmp_path / "validation.docx"
    _write_edit_doc(source)
    uploaded = await edit_api.upload_document(
        UploadFile(file=io.BytesIO(source.read_bytes()), filename="validation.docx")
    )
    doc_id = uploaded.id
    for operation in ("replace", "append", "insert"):
        with pytest.raises(HTTPException) as exc:
            await edit_api.edit_section(
                doc_id, SectionEditRequest(section_name="teaching_goals", operation=operation)
            )
        assert exc.value.status_code == 400
    with pytest.raises(HTTPException):
        await edit_api.edit_section(
            doc_id, SectionEditRequest(section_name="teaching_goals", operation="ai_modify")
        )
    with pytest.raises(HTTPException):
        await edit_api.edit_section(doc_id, SectionEditRequest(section_name="unknown", content="x"))
    with pytest.raises(HTTPException):
        await edit_api.add_section(doc_id, AddSectionRequest(section_name="teaching_goals"))
    with pytest.raises(HTTPException):
        await edit_api.ai_enhance_section(doc_id, AIEnhanceRequest(section_name="homework"))

    await api_db.execute(
        "UPDATE document_edits SET edit_history='[]', current_file_path=? WHERE id=?",
        (str(tmp_path / "missing.docx"), doc_id),
        commit=True,
    )
    with pytest.raises(HTTPException) as exc:
        await edit_api.undo_edit(doc_id)
    assert exc.value.status_code == 400
    with pytest.raises(HTTPException) as exc:
        await edit_api.save_document(doc_id)
    assert exc.value.status_code == 404
