import zipfile
import pathlib

import pytest
import pydantic
from docx import Document
from docx.oxml.ns import qn

from backend.config import settings
from backend.models import schemas
from backend.services import course_plan_renderer


def _chapters(count: int = 32, *, with_experiment_names: bool = True):
    return [
        schemas.ChapterInfo(
            lesson_number=index,
            topic=f"任务{index} Windows服务器安全配置",
            content_summary=f"掌握第{index}项服务器配置与安全验证方法",
            key_concepts=[f"配置{index}", "安全验证"],
            experiment_name=(
                f"安全配置实验{(index + 1) // 2}"
                if with_experiment_names and index % 2 == 1
                else ""
            ),
        )
        for index in range(1, count + 1)
    ]


def _all_text(path: str | pathlib.Path) -> str:
    document = Document(str(path))
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        ]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def _assert_preserved_parts(
    reference: pathlib.Path,
    generated: pathlib.Path,
    *,
    changed_parts: set[str] | None = None,
    added_parts: set[str] | None = None,
):
    changed = changed_parts or {"word/document.xml"}
    added = added_parts or set()
    with zipfile.ZipFile(reference) as source, zipfile.ZipFile(generated) as output:
        assert set(source.namelist()) | added == set(output.namelist())
        for name in source.namelist():
            if name not in changed:
                assert output.read(name) == source.read(name), name


def _teaching_week_rows(document: Document) -> list:
    """Collect week data rows across page tables in order."""
    rows = []
    for table in document.tables[:-1]:
        rows.extend(table.rows[2:])
    rows.extend(document.tables[-1].rows[2:-2])
    return rows


def _assert_paged_teaching_structure(path: pathlib.Path, week_count: int):
    document = Document(path)
    assert len(document.tables) >= 1
    for table in document.tables:
        # 每页一张表：2 列头行 + 数据行（末页另有合计与签字行）
        assert all(
            row._tr.trPr.find(qn("w:tblHeader")) is not None
            for row in table.rows[:2]
        )
        assert all(
            row._tr.trPr.find(qn("w:cantSplit")) is not None
            for row in table.rows
        )
    last_table = document.tables[-1]
    assert last_table.rows[-2].cells[0].paragraphs[0]._p.pPr.find(qn("w:keepNext")) is not None

    week_rows = _teaching_week_rows(document)
    assert [row.cells[0].text.strip() for row in week_rows] == [
        str(number) for number in range(1, week_count + 1)
    ]

    # 每页顶部都有标题与元数据段，页码为原位置的静态文字
    titles = [p.text for p in document.paragraphs if p.text == "云南林业职业技术学院教师授课计划表"]
    metas = [p.text for p in document.paragraphs if p.text.startswith("课名：")]
    page_count = len(document.tables)
    assert len(titles) == page_count
    assert len(metas) == page_count
    for page_no, meta in enumerate(metas, start=1):
        assert f"共 {page_count} 页" in meta
        assert f"第 {page_no} 页" in meta

    # 正文以 表格 + 1pt 占位段 + sectPr 结尾，避免空白页
    body_children = [
        child.tag.split("}")[1]
        for child in document.element.body.iterchildren()
    ]
    assert body_children[-3:] == ["tbl", "p", "sectPr"]


def test_fixed_course_plan_templates_are_valid():
    reports = course_plan_renderer.validate_all_course_plan_templates()

    assert {report["type"] for report in reports} == {
        "teaching_plan",
        "experiment_plan",
    }
    assert all(report["is_valid"] for report in reports)
    assert all(report["sha256"] == report["expected_sha256"] for report in reports)


def test_render_teaching_and_per_class_experiment_plans(tmp_path: pathlib.Path):
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)
    chapters = [chapter.model_dump() for chapter in _chapters()]
    for index, chapter in enumerate(chapters, start=1):
        chapter.update(
            key_points=f"准确完成第{index}项服务器配置并依据验证结果判断安全状态",
            difficult_points=f"综合分析第{index}项配置异常并选择合适方法完成故障排查",
            homework={
                "required": f"完成第{index}项配置验证记录",
                "optional": f"分析第{index}项安全加固方案",
            },
        )

    teaching_path = pathlib.Path(
        renderer.render_teaching_plan(
            batch_task_id="batch-1",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班", "24级信息安全技术应用2班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            total_hours=64,
            hours_per_lesson=2,
            start_week=1,
            chapters=chapters,
            location="慧心楼3516，实训楼204",
        )
    )
    experiment_paths = [
        pathlib.Path(path)
        for path in renderer.render_experiment_plans(
            batch_task_id="batch-1",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班", "24级信息安全技术应用2班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            plan_date="2026-02-25",
            first_class_date="2026-03-05",
            class_periods="3-4",
            hours_per_lesson=2,
            start_week=2,
            chapters=chapters,
            location="",
            class_schedules=[
                {
                    "class_name": "24级信息安全技术应用1班",
                    "weekday": 4,
                    "class_periods": "3-4",
                    "first_class_date": "2026-03-05",
                    "classroom": "慧心楼3516",
                },
                {
                    "class_name": "24级信息安全技术应用2班",
                    "weekday": 5,
                    "class_periods": "5-6",
                    "first_class_date": "2026-03-06",
                    "classroom": "实训楼204",
                },
            ],
        )
    ]

    assert teaching_path.is_file()
    assert len(experiment_paths) == 2
    teaching = Document(teaching_path)
    _assert_paged_teaching_structure(teaching_path, 16)
    assert teaching.tables[-1].rows[-2].cells[2].text == "32"
    assert teaching.tables[-1].rows[-2].cells[3].text == "32"
    assert teaching.tables[-1].rows[-2].cells[4].text == "64"
    teaching_text = _all_text(teaching_path)
    assert "24级信息安全技术应用1、2班" in teaching_text
    assert "机房、计算机" in teaching_text
    assert "任务32 Windows服务器安全配置" in teaching_text
    first_week = teaching.tables[0].rows[2].cells
    assert first_week[1].text == "任务1 Windows服务器安全配置\n任务2 Windows服务器安全配置"
    assert first_week[5].text == (
        "准确完成第1项服务器配置并依据验证结果判断安全状态\n"
        "准确完成第2项服务器配置并依据验证结果判断安全状态"
    )
    assert first_week[6].text == (
        "综合分析第1项配置异常并选择合适方法完成故障排查\n"
        "综合分析第2项配置异常并选择合适方法完成故障排查"
    )
    assert "…" not in first_week[5].text + first_week[6].text
    assert first_week[7].text == "机房、计算机"
    assert first_week[8].text == (
        "完成第1项配置验证记录\n分析第1项安全加固方案\n"
        "完成第2项配置验证记录\n分析第2项安全加固方案"
    )

    expectations = (
        ("24级信息安全技术应用1班", "四", "3-4", "05", "12", "慧心楼3516"),
        ("24级信息安全技术应用2班", "五", "5-6", "06", "13", "实训楼204"),
    )
    for output_index, (path, expectation) in enumerate(zip(
        experiment_paths,
        expectations,
    )):
        class_name, weekday, periods, first_day, second_day, classroom = expectation
        metadata_runs = Document(path).paragraphs[1].runs
        text = _all_text(path)
        assert class_name in text
        assert f"第 02 周  星期 {weekday}   第 {periods} 节     3 月  {first_day} 日" in text
        assert f"第 03 周  星期 {weekday}   第 {periods} 节     3 月  {second_day} 日" in text
        assert classroom in text
        assert "32    学时" in text
        assert all(metadata_runs[index].underline is True for index in (1, 3, 5, 7, 14, 16, 19, 21))
        experiment_table = Document(path).tables[0]
        assert len(experiment_table.rows) == 18
        assert all(
            row._tr.trPr.find(qn("w:cantSplit")) is not None
            for row in experiment_table.rows[1:]
        )
        assert experiment_table.rows[1].cells[3].text == classroom
        assert experiment_table.rows[1].cells[1].text == "安全配置实验1"
        assert "\n" not in experiment_table.rows[1].cells[1].text
        assert "…" not in experiment_table.rows[1].cells[1].text
        assert experiment_table.rows[1].cells[1]._tc.tcPr.find(qn("w:noWrap")) is not None
        assert all(
            any(cell.text.strip() for cell in row.cells)
            for row in experiment_table.rows[1:-1]
        )
        if output_index == 0:
            schedule_runs = experiment_table.rows[1].cells[2].paragraphs[0].runs
            assert [(run.text, run.underline) for run in schedule_runs] == [
                ("第", False),
                (" 02 ", True),
                ("周  星期", False),
                (" 四 ", True),
                ("  第", False),
                (" 3-4 ", True),
                ("节    ", False),
                (" 3 ", True),
                ("月 ", False),
                (" 05 ", True),
                ("日", False),
            ]

    _assert_preserved_parts(
        settings.teaching_plan_template_path,
        teaching_path,
        changed_parts={"word/document.xml"},
    )
    _assert_preserved_parts(settings.experiment_plan_template_path, experiment_paths[0])


def test_render_18_week_teaching_and_experiment_plans(tmp_path: pathlib.Path):
    chapters = _chapters(36)
    request = schemas.BatchTaskCreateRequest(
        course_name="Windows服务器安全配置",
        subject="信息安全技术",
        grade="2024级",
        template_id="yunlin-standard",
        total_hours=72,
        hours_per_lesson=2,
        chapters=chapters,
        start_week=1,
        class_ids=["class-1"],
        location="慧心楼3516",
        supplemental_artifacts=["teaching_plan", "experiment_plan"],
        academic_year="2025-2026",
        semester=2,
        teacher_name="李阳",
        plan_date="2026-02-25",
        first_class_date="2026-03-02",
        class_periods="3-4",
    )
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)

    teaching_path = pathlib.Path(
        renderer.render_teaching_plan(
            batch_task_id="batch-18-week",
            course_name=request.course_name,
            grade=request.grade,
            class_names=["24级信息安全技术应用1班"],
            academic_year=request.academic_year or "",
            semester=request.semester or 1,
            teacher_name=request.teacher_name or "",
            total_hours=request.total_hours,
            hours_per_lesson=request.hours_per_lesson,
            start_week=request.start_week,
            chapters=request.chapters,
            location=request.location,
        )
    )
    experiment_path = pathlib.Path(
        renderer.render_experiment_plans(
            batch_task_id="batch-18-week",
            course_name=request.course_name,
            grade=request.grade,
            class_names=["24级信息安全技术应用1班"],
            academic_year=request.academic_year or "",
            semester=request.semester or 1,
            teacher_name=request.teacher_name or "",
            plan_date=request.plan_date or "",
            first_class_date=request.first_class_date or "",
            class_periods=request.class_periods or "",
            hours_per_lesson=request.hours_per_lesson,
            start_week=request.start_week,
            chapters=request.chapters,
            location=request.location or "",
        )[0]
    )

    teaching = Document(teaching_path)
    _assert_paged_teaching_structure(teaching_path, 18)
    week_labels = [row.cells[0].text.strip() for row in _teaching_week_rows(teaching)]
    assert week_labels[-2:] == ["17", "18"]
    assert teaching.tables[-1].rows[-2].cells[1].text == "总课时"
    assert teaching.tables[-1].rows[-2].cells[4].text == "72"
    header_paragraphs = [
        p.text for p in teaching.paragraphs
        if p.text.startswith("课名：") or p.text.startswith("云南林业")
    ]
    header_text = "".join(header_paragraphs[:2])
    assert "网络安全" not in header_text
    assert "Windows服务器安全配置" in header_text
    assert "24级信息安全技术应用1班" in header_text

    experiment = Document(experiment_path)
    assert len(experiment.tables[0].rows) == 20
    assert experiment.tables[0].rows[18].cells[0].text == "实验18"
    assert "第 18 周" in experiment.tables[0].rows[18].cells[2].text
    assert "36    学时" in experiment.tables[0].rows[-1].cells[2].text

    _assert_preserved_parts(
        settings.teaching_plan_template_path,
        teaching_path,
        changed_parts={"word/document.xml"},
    )
    _assert_preserved_parts(settings.experiment_plan_template_path, experiment_path)


def test_teaching_plan_homework_is_brief_and_varied(tmp_path: pathlib.Path):
    chapters = [chapter.model_dump() for chapter in _chapters(6)]
    long_homework = [
        # 第 1 周：报告类 → 撰写实验报告
        {"required": "完成实验项目《异常处理与自定义异常》的代码和实验报告；另编写一个用户登录模拟程序并提交运行截图", "optional": ""},
        "撰写实验报告并提交平台，完成课后拓展任务与自查清单，思考异常处理与日志记录的关系",
        # 第 2 周：练习类 → 完成课后练习
        {"required": "认真完成课后练习第五题到第十二题，整理错题并进行巩固训练后提交到学习平台", "optional": ""},
        {"required": "完成教材配套习题册中本节对应的全部题目，核对答案后将错题订正提交", "optional": ""},
        # 第 3 周：实训类 → 上机实操评估
        {"required": "在实训室完成交换机VLAN划分与端口配置，记录配置过程并截图提交", "optional": ""},
        {"required": "分组完成路由器静态路由配置实训项目，互相检查并演示验收结果", "optional": ""},
    ]
    for chapter, homework in zip(chapters, long_homework):
        chapter["homework"] = homework
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)

    path = pathlib.Path(
        renderer.render_teaching_plan(
            batch_task_id="batch-brief-homework",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            total_hours=12,
            hours_per_lesson=2,
            start_week=1,
            chapters=chapters,
            location="慧心楼3516",
        )
    )

    rows = _teaching_week_rows(Document(path))
    homework_cells = [rows[week].cells[8].text for week in range(3)]
    assert homework_cells == ["撰写实验报告", "完成课后练习", "上机实操评估"]
    assert all(
        len(line) <= course_plan_renderer.HOMEWORK_MAX_CHARS
        for cell in homework_cells
        for line in cell.splitlines()
    )


def test_teaching_plan_points_are_brief(tmp_path: pathlib.Path):
    chapters = [chapter.model_dump() for chapter in _chapters(2)]
    chapters[0]["key_points"] = (
        "掌握raise语句主动抛出异常的方法；掌握assert断言的语法与适用场景；"
        "掌握自定义异常类的定义及异常类型继承机制"
    )
    chapters[0]["difficult_points"] = "依据业务逻辑设计合理的自定义异常类并正确处理异常继承关系，区分raise和assert的使用场景"
    chapters[1]["key_points"] = "掌握try-except语句结构与多异常捕获"
    chapters[1]["difficult_points"] = "理解异常传播流程"
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)

    path = pathlib.Path(
        renderer.render_teaching_plan(
            batch_task_id="batch-brief-points",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            total_hours=4,
            hours_per_lesson=2,
            start_week=1,
            chapters=chapters,
            location="慧心楼3516",
        )
    )

    first_week = _teaching_week_rows(Document(path))[0].cells
    # 多要点句只保留装得下的前几条（16 + 1 + 15 = 32 > 25，仅保留第一条）
    assert first_week[5].text == (
        "掌握raise语句主动抛出异常的方法\n掌握try-except语句结构与多异常捕获"
    )
    # 无分号的长难点硬截断到 25 字
    difficult_lines = first_week[6].text.splitlines()
    assert difficult_lines == [
        "依据业务逻辑设计合理的自定义异常类并正确处理异常继",
        "理解异常传播流程",
    ]
    assert all(
        len(line) <= course_plan_renderer.POINTS_MAX_CHARS
        for cell in (first_week[5], first_week[6])
        for line in cell.text.splitlines()
    )


def test_batch_plan_request_rejects_more_than_18_weeks():
    with pytest.raises(pydantic.ValidationError, match="最多支持 18 周"):
        schemas.BatchTaskCreateRequest(
            course_name="网络安全",
            subject="信息安全技术",
            grade="2024级",
            template_id="yunlin-standard",
            total_hours=76,
            chapters=_chapters(38),
            class_ids=["class-1"],
            supplemental_artifacts=["teaching_plan"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
        )


def test_explicit_experiment_names_filter_non_experiment_weeks(tmp_path: pathlib.Path):
    chapters = _chapters(8, with_experiment_names=False)
    chapters[0].experiment_name = "认识Windows服务"
    chapters[4].experiment_name = "配置IIS安全功能"
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)

    output = pathlib.Path(
        renderer.render_experiment_plans(
            batch_task_id="batch-2",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            plan_date="2026-02-25",
            first_class_date="2026-03-05",
            class_periods="3-4",
            hours_per_lesson=2,
            start_week=2,
            chapters=chapters,
            location="慧心楼3516",
        )[0]
    )
    document = Document(output)

    assert len(document.tables[0].rows) == 4
    assert document.tables[0].rows[1].cells[1].text == "认识Windows服务"
    assert document.tables[0].rows[2].cells[1].text == "配置IIS安全功能"
    assert "第 04 周" in document.tables[0].rows[2].cells[2].text
    assert "4    学时" in document.tables[0].rows[-1].cells[2].text


def test_experiment_plan_rejects_names_that_would_wrap(tmp_path: pathlib.Path):
    chapters = _chapters(2, with_experiment_names=False)
    chapters[0].experiment_name = "这是一个明显超过固定模板单行容量的实验项目名称"
    renderer = course_plan_renderer.CoursePlanRenderer(tmp_path)

    with pytest.raises(ValueError, match="不能超过 18 个字符"):
        renderer.render_experiment_plans(
            batch_task_id="batch-invalid-name",
            course_name="Windows服务器安全配置",
            grade="24级",
            class_names=["24级信息安全技术应用1班"],
            academic_year="2025-2026",
            semester=2,
            teacher_name="李阳",
            plan_date="2026-02-25",
            first_class_date="2026-03-05",
            class_periods="3-4",
            hours_per_lesson=2,
            start_week=2,
            chapters=chapters,
            location="慧心楼3516",
        )


def test_batch_plan_request_requires_fixed_schedule_fields():
    with pytest.raises(pydantic.ValidationError, match="学年"):
        schemas.BatchTaskCreateRequest(
            course_name="网络安全",
            subject="信息安全技术",
            grade="2024级",
            template_id="yunlin-standard",
            total_hours=4,
            chapters=_chapters(2),
            class_ids=["class-1"],
            supplemental_artifacts=["teaching_plan"],
        )

    request = schemas.BatchTaskCreateRequest(
        course_name="网络安全",
        subject="信息安全技术",
        grade="2024级",
        template_id="yunlin-standard",
        total_hours=4,
        chapters=_chapters(2),
        class_ids=["class-1"],
        supplemental_artifacts=["teaching_plan"],
        academic_year="2025-2026",
        semester=2,
        teacher_name="李阳",
    )
    assert request.plan_date is None
