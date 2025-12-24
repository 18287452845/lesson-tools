"""
Lesson plan parser service for analyzing existing Word documents.
This service intelligently identifies sections in existing lesson plans.
"""
import re
from pathlib import Path
from typing import Optional, Dict, List, Any
from dataclasses import dataclass, field
from docx import Document

from ..models.schemas import ParsedSection, SectionLocation


@dataclass
class SectionInfo:
    """Information about a parsed section."""
    name: str
    found: bool
    content: Optional[str]
    start_para_idx: Optional[int] = None
    end_para_idx: Optional[int] = None
    in_table: bool = False
    table_idx: Optional[int] = None
    cell_location: Optional[tuple] = None
    start_table_row: Optional[int] = None
    start_table_col: Optional[int] = None


class LessonPlanParser:
    """
    Parser for existing lesson plan Word documents.

    Intelligently identifies common lesson plan sections using:
    1. Keyword-based header matching
    2. Document structure analysis (headings, tables)
    3. Content pattern recognition

    Supports Chinese and English section headers.
    """

    # Section name patterns - regex patterns for matching section headers
    SECTION_PATTERNS = {
        "teaching_goals": [
            r"教学目标\s*[：:]*",
            r"三维目标\s*[：:]*",
            r"学习目标\s*[：:]*",
            r"课程目标\s*[：:]*",
            r"Teaching\s+Goals?\s*[：:]*",
            r"Learning\s+Objectives?\s*[：:]*",
        ],
        "key_points": [
            r"教学重点\s*[：:]*",
            r"重点\s*[：:]*",
            r"本课重点\s*[：:]*",
            r"Key\s+Points?\s*[：:]*",
        ],
        "difficult_points": [
            r"教学难点\s*[：:]*",
            r"难点\s*[：:]*",
            r"本课难点\s*[：:]*",
            r"Difficult\s+Points?\s*[：:]*",
        ],
        "teaching_tools": [
            r"教具?\s*(?:准备)?\s*[：:]*",
            r"学具\s*(?:准备)?\s*[：:]*",
            r"教学准备\s*[：:]*",
            r"课前准备\s*[：:]*",
            r"Teaching\s+Aids?\s*[：:]*",
            r"Materials?\s*(?:Needed)?\s*[：:]*",
        ],
        "teaching_methods": [
            r"教学方法\s*[：:]*",
            r"授课方法\s*[：:]*",
            r"Teaching\s+Methods?\s*[：:]*",
        ],
        "teaching_process": [
            r"教学过程\s*[：:]*",
            r"教学环节\s*[：:]*",
            r"教学步骤\s*[：:]*",
            r"教学流程\s*[：:]*",
            r"Teaching\s+Process\s*[：:]*",
            r"Teaching\s+Procedures?\s*[：:]*",
        ],
        "student_analysis": [
            r"学情分析\s*[：:]*",
            r"学生分析\s*[：:]*",
            r"Student\s+Analysis\s*[：:]*",
            r"Learner\s+Analysis\s*[：:]*",
        ],
        "textbook_analysis": [
            r"教材分析\s*[：:]*",
            r"Textbook\s+Analysis\s*[：:]*",
        ],
        "homework": [
            r"作业\s*(?:布置|设计)?\s*[：:]*",
            r"课后作业\s*[：:]*",
            r"课后练习\s*[：:]*",
            r"布置作业\s*[：:]*",
            r"Homework\s*[：:]*",
            r"Assignments?\s*[：:]*",
        ],
        "reflection": [
            r"教学反思\s*[：:]*",
            r"课后反思\s*[：:]*",
            r"教后[记感]\s*[：:]*",
            r"Reflection\s*[：:]*",
        ],
        "blackboard_design": [
            r"板书\s*(?:设计)?\s*[：:]*",
            r"Blackboard\s+Design\s*[：:]*",
        ],
    }

    # Process stage patterns (for detailed process parsing)
    PROCESS_STAGE_PATTERNS = [
        r"(?:导入|引入)\s*(?:新课)?\s*[：:]*",
        r"探究\s*(?:新知)\s*[：:]*",
        r"(?:新授|讲授)\s*(?:新知)?\s*[：:]*",
        r"(?:巩固|练习)\s*(?:练习)?\s*[：:]*",
        r"课堂\s*(?:小结|总结)\s*[：:]*",
        r"拓展\s*(?:延伸)?\s*[：:]*",
    ]

    def __init__(self, doc_path: str | Path):
        """Initialize the parser with a document path."""
        self.doc_path = Path(doc_path)
        self.doc: Optional[Document] = None
        self.sections: Dict[str, SectionInfo] = {}

    def parse(self) -> Dict[str, SectionInfo]:
        """
        Parse the entire document and identify all sections.

        Returns:
            Dictionary mapping section names to SectionInfo objects
        """
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {self.doc_path}")

        self.doc = Document(str(self.doc_path))
        self.sections = {}

        # Parse tables first (table-based layouts are common in lesson plans)
        self._parse_tables()

        # Parse paragraphs (may override or supplement table findings)
        self._parse_paragraphs()

        # Mark missing sections
        self._mark_missing_sections()

        return self.sections

    def _parse_tables(self) -> None:
        """Parse content in table structures."""
        for table_idx, table in enumerate(self.doc.tables):
            self._parse_single_table(table, table_idx)

    def _parse_single_table(self, table, table_idx: int) -> None:
        """Parse a single table for section content."""
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        # Common layout: first column is headers, second is content
        if cols >= 2:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) < 2:
                    continue

                header_cell = row.cells[0]
                content_cell = row.cells[1]

                header_text = header_cell.text.strip()
                content_text = content_cell.text.strip()

                matched_section = self._match_section_header(header_text)
                if matched_section:
                    self.sections[matched_section] = SectionInfo(
                        name=matched_section,
                        found=True,
                        content=content_text if content_text else None,
                        start_para_idx=None,
                        end_para_idx=None,
                        in_table=True,
                        table_idx=table_idx,
                        cell_location=(row_idx, 1),
                        start_table_row=row_idx,
                        start_table_col=0,
                    )

        # Also check all cells for section headers
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                matched_section = self._match_section_header(cell_text)

                if matched_section and matched_section not in self.sections:
                    # Try to find content in adjacent cells
                    content = self._extract_adjacent_content(table, row_idx, col_idx)

                    # Extract content after header in the same cell
                    header_content = self._extract_content_after_header(
                        cell_text, matched_section
                    )
                    if header_content:
                        content = header_content

                    self.sections[matched_section] = SectionInfo(
                        name=matched_section,
                        found=True,
                        content=content if content else None,
                        start_para_idx=None,
                        end_para_idx=None,
                        in_table=True,
                        table_idx=table_idx,
                        cell_location=(row_idx, col_idx),
                        start_table_row=row_idx,
                        start_table_col=col_idx,
                    )

    def _extract_adjacent_content(
        self, table, row_idx: int, col_idx: int
    ) -> Optional[str]:
        """Extract content from cells adjacent to a header cell."""
        content_parts = []

        # Check right cell
        if col_idx + 1 < len(table.rows[row_idx].cells):
            right_cell = table.rows[row_idx].cells[col_idx + 1]
            if right_cell.text.strip():
                content_parts.append(right_cell.text.strip())

        # Check cells below (for merged content scenarios)
        if row_idx + 1 < len(table.rows):
            next_row = table.rows[row_idx + 1]
            # Look for cells that span the width (likely content)
            for cell in next_row.cells:
                if cell.text.strip() and len(cell.text) > 20:  # Substantial content
                    content_parts.append(cell.text.strip())

        return "\n".join(content_parts) if content_parts else None

    def _parse_paragraphs(self) -> None:
        """Parse paragraph-based document structure."""
        if self.doc is None:
            return

        current_section = None
        current_content: List[str] = []
        start_idx = None

        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check if this is a section header
            matched_section = self._match_section_header(text)

            if matched_section:
                # Save the previous section if it exists and not already found in tables
                if (
                    current_section
                    and current_section not in self.sections
                    and current_content
                ):
                    self.sections[current_section] = SectionInfo(
                        name=current_section,
                        found=True,
                        content="\n".join(current_content),
                        start_para_idx=start_idx,
                        end_para_idx=idx - 1,
                        in_table=False,
                    )

                # Start tracking the new section
                current_section = matched_section
                current_content = []
                start_idx = idx

                # Extract content immediately after the header (e.g., "重点：xxx")
                header_content = self._extract_content_after_header(
                    text, matched_section
                )
                if header_content:
                    current_content.append(header_content)

            elif current_section and current_section not in self.sections:
                # Continue collecting content for current section
                current_content.append(text)

        # Save the last section
        if (
            current_section
            and current_section not in self.sections
            and current_content
        ):
            self.sections[current_section] = SectionInfo(
                name=current_section,
                found=True,
                content="\n".join(current_content),
                start_para_idx=start_idx,
                end_para_idx=len(self.doc.paragraphs) - 1,
                in_table=False,
            )

    def _match_section_header(self, text: str) -> Optional[str]:
        """
        Match text against section header patterns.

        Args:
            text: The text to match

        Returns:
            The section name if matched, None otherwise
        """
        for section_name, patterns in self.SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return section_name
        return None

    def _extract_content_after_header(
        self, text: str, section_name: str
    ) -> Optional[str]:
        """
        Extract content that appears after a section header on the same line.
        e.g., "教学重点：理解分数的概念" -> "理解分数的概念"
        """
        patterns = self.SECTION_PATTERNS.get(section_name, [])
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                content = text[match.end() :].strip()
                return content if content else None
        return None

    def _mark_missing_sections(self) -> None:
        """Mark sections that were not found in the document."""
        all_sections = set(self.SECTION_PATTERNS.keys())
        found_sections = set(self.sections.keys())
        missing_sections = all_sections - found_sections

        for section in missing_sections:
            self.sections[section] = SectionInfo(
                name=section,
                found=False,
                content=None,
                start_para_idx=None,
                end_para_idx=None,
            )

    def get_parsed_sections(self) -> Dict[str, ParsedSection]:
        """
        Get parsed sections in the API response format.

        Returns:
            Dictionary of ParsedSection objects
        """
        result = {}
        for section_name, section_info in self.sections.items():
            location = None
            if section_info.found:
                location = SectionLocation(
                    in_table=section_info.in_table,
                    table_idx=section_info.table_idx,
                    cell_location=section_info.cell_location,
                    start_para_idx=section_info.start_para_idx,
                    end_para_idx=section_info.end_para_idx,
                )

            result[section_name] = ParsedSection(
                name=section_name,
                found=section_info.found,
                content=section_info.content,
                location=location,
            )
        return result

    def get_section_content(self, section_name: str) -> Optional[str]:
        """
        Get the content of a specific section.

        Args:
            section_name: Name of the section

        Returns:
            Section content or None if not found
        """
        section = self.sections.get(section_name)
        return section.content if section and section.found else None

    def is_section_found(self, section_name: str) -> bool:
        """
        Check if a section was found in the document.

        Args:
            section_name: Name of the section

        Returns:
            True if section was found
        """
        section = self.sections.get(section_name)
        return section.found if section else False

    def get_document_summary(self) -> Dict[str, Any]:
        """
        Get a summary of the parsed document.

        Returns:
            Dictionary with document summary
        """
        if self.doc is None:
            return {}

        found_count = sum(1 for s in self.sections.values() if s.found)
        total_count = len(self.sections)

        return {
            "filename": self.doc_path.name,
            "paragraph_count": len(self.doc.paragraphs),
            "table_count": len(self.doc.tables),
            "sections_found": found_count,
            "sections_total": total_count,
            "sections": {name: info.found for name, info in self.sections.items()},
        }

    def parse_teaching_process_stages(self) -> List[Dict[str, Any]]:
        """
        Parse the teaching process into detailed stages if possible.

        Returns:
            List of stage dictionaries
        """
        process_content = self.get_section_content("teaching_process")
        if not process_content:
            return []

        stages = []
        lines = process_content.split("\n")
        current_stage = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line matches a stage pattern
            for pattern in self.PROCESS_STAGE_PATTERNS:
                if re.search(pattern, line):
                    if current_stage:
                        stages.append(current_stage)

                    stage_name = re.sub(pattern, "", line).strip()
                    current_stage = {
                        "stage": stage_name or line.split(":")[0],
                        "content": [],
                    }
                    break
            else:
                # Content line
                if current_stage is not None:
                    current_stage["content"].append(line)
                elif ":" in line or "：" in line:
                    # Try to create a stage from header-like content
                    parts = re.split(r"[:：]", line, 1)
                    current_stage = {"stage": parts[0], "content": [parts[1]] if len(parts) > 1 else []}

        if current_stage:
            stages.append(current_stage)

        return stages


def parse_document(doc_path: str) -> Dict[str, ParsedSection]:
    """
    Convenience function to parse a lesson plan document.

    Args:
        doc_path: Path to the document

    Returns:
        Dictionary of ParsedSection objects
    """
    parser = LessonPlanParser(doc_path)
    parser.parse()
    return parser.get_parsed_sections()


def get_document_summary(doc_path: str) -> Dict[str, Any]:
    """
    Convenience function to get a document summary.

    Args:
        doc_path: Path to the document

    Returns:
        Document summary dictionary
    """
    parser = LessonPlanParser(doc_path)
    parser.parse()
    return parser.get_document_summary()
