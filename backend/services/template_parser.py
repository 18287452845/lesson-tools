"""
Template parser service for analyzing Word document templates.
"""
import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
from docx import Document
from docx.oxml import parse_xml

from ..models.schemas import FieldConfig


@dataclass
class TemplateField:
    """A field found in a template."""
    name: str
    raw_placeholder: str
    field_type: str  # simple, loop_start, loop_end, conditional
    loop_variable: Optional[str] = None
    line_number: Optional[int] = None


class TemplateParser:
    """
    Parser for Word document templates.

    Supports:
    - Simple variables: {{variable_name}}
    - Loop structures: {% for item in items %} ... {% endfor %}
    - Conditionals: {% if condition %} ... {% endif %}
    """

    # Pattern for Jinja2-style variables
    VARIABLE_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")

    # Pattern for loop structures
    LOOP_START_PATTERN = re.compile(r"\{%\s*for\s+(\w+)\s+in\s+(\w+)\s*%\}")
    LOOP_END_PATTERN = re.compile(r"\{%\s*endfor\s*%\}")

    # Pattern for conditionals
    IF_START_PATTERN = re.compile(r"\{%\s*if\s+(\w+)\s*%\}")
    IF_END_PATTERN = re.compile(r"\{%\s*endif\s*%\}")

    # Standard field mappings
    STANDARD_FIELDS = {
        "subject": {"display": "学科", "type": "text", "required": True},
        "grade": {"display": "年级", "type": "text", "required": True},
        "topic": {"display": "课题", "type": "text", "required": True},
        "duration": {"display": "课时", "type": "text", "required": True},
        "teaching_goals": {"display": "教学目标", "type": "textarea", "required": True},
        "key_points": {"display": "教学重点", "type": "textarea", "required": True},
        "difficult_points": {"display": "教学难点", "type": "textarea", "required": True},
        "teaching_tools": {"display": "教具准备", "type": "textarea", "required": False},
        "teaching_methods": {"display": "教学方法", "type": "textarea", "required": False},
        "student_analysis": {"display": "学情分析", "type": "textarea", "required": False},
        "textbook_analysis": {"display": "教材分析", "type": "textarea", "required": False},
        "homework": {"display": "作业布置", "type": "textarea", "required": False},
        "blackboard_design": {"display": "板书设计", "type": "textarea", "required": False},
        "reflection": {"display": "教学反思", "type": "textarea", "required": False},
        "teaching_steps": {"display": "教学过程", "type": "json", "required": True},
    }

    def __init__(self, template_path: str | Path):
        """Initialize the parser with a template file path."""
        self.template_path = Path(template_path)
        self.doc = None
        self.fields: List[TemplateField] = []
        self.field_configs: List[FieldConfig] = []

    def parse(self) -> List[FieldConfig]:
        """
        Parse the template and extract field configurations.

        Returns:
            List of FieldConfig objects
        """
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        self.doc = Document(str(self.template_path))
        self.fields = []

        # Parse paragraphs
        for para_idx, para in enumerate(self.doc.paragraphs):
            self._parse_text(para.text, line_number=para_idx)

        # Parse tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    self._parse_text(cell.text, line_number=f"T{table_idx}R{row_idx}C{cell_idx}")

        # Convert to field configs
        self.field_configs = self._create_field_configs()

        return self.field_configs

    def _parse_text(self, text: str, line_number: Any = None) -> None:
        """Parse text and extract template fields."""
        # Find all variables
        for match in self.VARIABLE_PATTERN.finditer(text):
            var_name = match.group(1)
            placeholder = match.group(0)
            self.fields.append(
                TemplateField(
                    name=var_name,
                    raw_placeholder=placeholder,
                    field_type="simple",
                    line_number=line_number,
                )
            )

        # Find loop structures
        if self.LOOP_START_PATTERN.search(text):
            match = self.LOOP_START_PATTERN.search(text)
            loop_var = match.group(1)
            loop_iter = match.group(2)
            self.fields.append(
                TemplateField(
                    name=loop_iter,
                    raw_placeholder=match.group(0),
                    field_type="loop_start",
                    loop_variable=loop_var,
                    line_number=line_number,
                )
            )

        if self.LOOP_END_PATTERN.search(text):
            self.fields.append(
                TemplateField(
                    name="",
                    raw_placeholder="{% endfor %}",
                    field_type="loop_end",
                    line_number=line_number,
                )
            )

        # Find conditionals
        if self.IF_START_PATTERN.search(text):
            match = self.IF_START_PATTERN.search(text)
            condition_var = match.group(1)
            self.fields.append(
                TemplateField(
                    name=condition_var,
                    raw_placeholder=match.group(0),
                    field_type="conditional",
                    line_number=line_number,
                )
            )

        if self.IF_END_PATTERN.search(text):
            self.fields.append(
                TemplateField(
                    name="",
                    raw_placeholder="{% endif %}",
                    field_type="if_end",
                    line_number=line_number,
                )
            )

    def _create_field_configs(self) -> List[FieldConfig]:
        """Create FieldConfig objects from parsed fields."""
        configs = []
        seen_fields = set()

        for field in self.fields:
            # Skip structural elements
            if field.field_type in ("loop_end", "if_end"):
                continue

            # Skip duplicates
            if field.name in seen_fields:
                continue

            seen_fields.add(field.name)

            # Check if it's a standard field
            if field.name in self.STANDARD_FIELDS:
                std_field = self.STANDARD_FIELDS[field.name]
                configs.append(
                    FieldConfig(
                        name=field.name,
                        display_name=std_field["display"],
                        description=f"输入{std_field['display']}",
                        required=std_field["required"],
                        field_type=std_field["type"],
                    )
                )
            else:
                # Determine field type from structure
                field_type = "text"
                if field.field_type == "loop_start":
                    field_type = "array"
                elif field.name.endswith("_analysis") or field.name.endswith("_content"):
                    field_type = "textarea"

                configs.append(
                    FieldConfig(
                        name=field.name,
                        display_name=field.name.replace("_", " ").title(),
                        description=f"输入{field.name}",
                        required=False,
                        field_type=field_type,
                    )
                )

        return configs

    def get_template_structure(self) -> Dict[str, Any]:
        """
        Get the template structure summary.

        Returns:
            Dictionary with template structure information
        """
        if self.doc is None:
            self.parse()

        return {
            "paragraph_count": len(self.doc.paragraphs),
            "table_count": len(self.doc.tables),
            "field_count": len(self.field_configs),
            "fields": [f.model_dump() for f in self.field_configs],
            "has_loops": any(f.field_type == "loop_start" for f in self.fields),
            "has_conditionals": any(f.field_type == "conditional" for f in self.fields),
        }

    def validate_template(self) -> tuple[bool, List[str]]:
        """
        Validate the template structure.

        Returns:
            Tuple of (is_valid, error_messages)
        """
        errors = []

        if self.doc is None:
            try:
                self.parse()
            except Exception as e:
                return False, [f"Failed to parse template: {e}"]

        # Check for matching loop tags
        loop_starts = sum(1 for f in self.fields if f.field_type == "loop_start")
        loop_ends = sum(1 for f in self.fields if f.field_type == "loop_end")

        if loop_starts != loop_ends:
            errors.append(f"Mismatched loop tags: {loop_starts} start, {loop_ends} end")

        # Check for matching if tags
        if_starts = sum(1 for f in self.fields if f.field_type == "conditional")
        if_ends = sum(1 for f in self.fields if f.field_type == "if_end")

        if if_starts != if_ends:
            errors.append(f"Mismatched if tags: {if_starts} start, {if_ends} end")

        # Check for required fields
        found_fields = {f.name for f in self.fields}
        required_fields = {"subject", "grade", "topic", "duration"}
        missing_fields = required_fields - found_fields

        if missing_fields:
            errors.append(f"Missing required fields: {', '.join(missing_fields)}")

        return len(errors) == 0, errors


def parse_template_file(template_path: str) -> List[FieldConfig]:
    """
    Convenience function to parse a template file.

    Args:
        template_path: Path to the template file

    Returns:
        List of FieldConfig objects
    """
    parser = TemplateParser(template_path)
    return parser.parse()


def validate_template_file(template_path: str) -> tuple[bool, List[str]]:
    """
    Convenience function to validate a template file.

    Args:
        template_path: Path to the template file

    Returns:
        Tuple of (is_valid, error_messages)
    """
    parser = TemplateParser(template_path)
    return parser.validate_template()
