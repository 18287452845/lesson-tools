"""
Word 文档转换器

实现 .docx ↔ HTML 双向转换，并保护 Jinja2 模板语法
"""
import os
import tempfile
from typing import Dict, Any, Optional
from pathlib import Path

import mammoth
from docx import Document
from htmldocx import HtmlToDocx
from bs4 import BeautifulSoup

from .jinja_protector import protect_jinja_syntax, restore_jinja_syntax


class DocxConverter:
    """Word 文档与 HTML 的转换器"""

    # Mammoth 样式映射配置
    STYLE_MAP = """
        p[style-name='Heading 1'] => h1.heading-1
        p[style-name='Heading 2'] => h2.heading-2
        p[style-name='Heading 3'] => h3.heading-3
        p[style-name='标题 1'] => h1.heading-1
        p[style-name='标题 2'] => h2.heading-2
        p[style-name='标题 3'] => h3.heading-3
        table => table.word-table[data-preserve='true']
        r[style-name='Strong'] => strong
        r[style-name='加粗'] => strong
        r[style-name='Emphasis'] => em
        r[style-name='斜体'] => em
    """

    def __init__(self):
        self.html_to_docx = HtmlToDocx()

    def to_html(self, docx_path: str) -> Dict[str, Any]:
        """
        将 .docx 文件转换为 HTML，并保护 Jinja2 语法

        Args:
            docx_path: .docx 文件路径

        Returns:
            {
                'html': str,           # HTML 内容
                'messages': List,      # 转换消息
                'metadata': Dict,      # 文档元数据
                'original_path': str   # 原始文件路径
            }
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")

        # 使用 mammoth 转换
        with open(docx_path, 'rb') as f:
            result = mammoth.convert_to_html(
                f,
                style_map=self.STYLE_MAP,
                convert_image=mammoth.images.inline(self._base64_image)
            )

        html = result.value
        messages = result.messages

        # 保护 Jinja2 语法
        html = protect_jinja_syntax(html)

        # 提取元数据
        metadata = self._extract_metadata(docx_path)

        # 增强 HTML（添加必要的包装和样式）
        html = self._enhance_html(html)

        return {
            'html': html,
            'messages': [str(msg) for msg in messages],
            'metadata': metadata,
            'original_path': docx_path
        }

    def to_docx(
        self,
        html: str,
        output_path: str,
        original_docx_path: Optional[str] = None
    ) -> str:
        """
        将 HTML 转换为 .docx 文件，并恢复 Jinja2 语法

        Args:
            html: HTML 内容
            output_path: 输出文件路径
            original_docx_path: 原始 .docx 文件路径（用于复制样式）

        Returns:
            输出文件路径
        """
        # 恢复 Jinja2 语法
        html = restore_jinja_syntax(html)

        # 清理 HTML（移除编辑器添加的额外标记）
        html = self._clean_html(html)

        # 创建 Word 文档
        if original_docx_path and os.path.exists(original_docx_path):
            # 基于原始文档创建（保留样式）
            doc = Document(original_docx_path)
            # 清空内容但保留样式
            for element in doc.element.body:
                element.getparent().remove(element)
        else:
            # 创建新文档
            doc = Document()

        # 将 HTML 转换为 Word
        self.html_to_docx.add_html_to_document(html, doc)

        # 保存文档
        doc.save(output_path)

        return output_path

    def preview_html(self, html: str, sample_data: Dict[str, Any]) -> str:
        """
        生成预览 HTML（用 示例数据渲染 Jinja2）

        Args:
            html: 包含 Jinja2 的 HTML
            sample_data: 示例数据

        Returns:
            渲染后的 HTML
        """
        # 恢复 Jinja2 语法
        jinja_html = restore_jinja_syntax(html)

        # 使用 nunjucks 或简单替换进行预览
        # 这里先实现简单版本
        preview = jinja_html

        # 简单替换变量
        for key, value in sample_data.items():
            if isinstance(value, str):
                preview = preview.replace(f'{{{{{key}}}}}', value)
            elif isinstance(value, (list, dict)):
                # 复杂类型暂时显示为 JSON
                import json
                preview = preview.replace(
                    f'{{{{{key}}}}}',
                    json.dumps(value, ensure_ascii=False)
                )

        return preview

    def _extract_metadata(self, docx_path: str) -> Dict[str, Any]:
        """提取文档元数据"""
        try:
            doc = Document(docx_path)
            core_props = doc.core_properties

            metadata = {
                'title': core_props.title or '',
                'subject': core_props.subject or '',
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'file_size': os.path.getsize(docx_path),
                'file_name': os.path.basename(docx_path),
            }

            # 统计信息
            metadata['paragraphs_count'] = len(doc.paragraphs)
            metadata['tables_count'] = len(doc.tables)

            return metadata
        except Exception as e:
            return {'error': str(e)}

    def _enhance_html(self, html: str) -> str:
        """
        增强 HTML 内容

        添加必要的包装、样式类等
        """
        soup = BeautifulSoup(html, 'html.parser')

        # 为表格添加类
        for table in soup.find_all('table'):
            if not table.get('class'):
                table['class'] = ['word-table']
            # 确保有 data-preserve 属性
            if 'data-preserve' not in table.attrs:
                table['data-preserve'] = 'true'

        # 为 Jinja2 占位符添加样式提示
        for span in soup.find_all('span', class_='jinja-placeholder'):
            # 确保不可编辑
            span['contenteditable'] = 'false'
            # 添加视觉提示
            if 'jinja-variable' in span.get('class', []):
                span['style'] = (
                    'background-color: #e6f7ff; '
                    'border: 1px solid #1890ff; '
                    'border-radius: 3px; '
                    'padding: 2px 6px; '
                    'margin: 0 2px; '
                    'font-family: monospace; '
                    'cursor: pointer;'
                )

        return str(soup)

    def _clean_html(self, html: str) -> str:
        """
        清理 HTML，移除编辑器添加的额外标记

        保留 Word 需要的结构，移除不必要的属性
        """
        soup = BeautifulSoup(html, 'html.parser')

        # 移除编辑器相关的属性
        for tag in soup.find_all():
            # 保留的属性
            keep_attrs = ['class', 'style', 'colspan', 'rowspan', 'data-preserve']

            # 移除其他属性
            attrs_to_remove = [
                attr for attr in tag.attrs
                if attr not in keep_attrs and not attr.startswith('data-jinja')
            ]

            for attr in attrs_to_remove:
                del tag[attr]

        # 移除空段落（但保留有 Jinja2 占位符的）
        for p in soup.find_all('p'):
            if not p.get_text(strip=True) and not p.find('span', class_='jinja-placeholder'):
                p.decompose()

        return str(soup)

    def _base64_image(self, image):
        """将图片转换为 base64 编码（用于内联图片）"""
        with image.open() as image_bytes:
            encoded_string = image_bytes.read()
            import base64
            return {
                "src": f"data:{image.content_type};base64,{base64.b64encode(encoded_string).decode()}"
            }

    def validate_docx(self, docx_path: str) -> tuple[bool, list[str]]:
        """
        验证 .docx 文件是否可以安全转换

        Args:
            docx_path: .docx 文件路径

        Returns:
            (是否有效, 错误/警告列表)
        """
        errors = []

        # 检查文件存在
        if not os.path.exists(docx_path):
            errors.append(f"文件不存在: {docx_path}")
            return False, errors

        # 检查文件扩展名
        if not docx_path.lower().endswith('.docx'):
            errors.append("文件必须是 .docx 格式")
            return False, errors

        # 尝试打开文档
        try:
            doc = Document(docx_path)

            # 检查是否有内容
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                errors.append("文档为空")

            # 检查文档大小（超过 10MB 警告）
            file_size = os.path.getsize(docx_path)
            if file_size > 10 * 1024 * 1024:
                errors.append(f"警告：文件较大 ({file_size / 1024 / 1024:.2f} MB)，转换可能较慢")

            # 检查复杂元素
            has_images = any(
                run._element.xpath('.//pic:pic', namespaces={
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                })
                for para in doc.paragraphs
                for run in para.runs
            )

            if has_images:
                errors.append("警告：文档包含图片，部分格式可能无法完美保留")

        except Exception as e:
            errors.append(f"无法打开文档: {str(e)}")
            return False, errors

        # 如果只有警告，仍然返回 True
        has_real_errors = any(
            not msg.startswith('警告：')
            for msg in errors
        )

        return not has_real_errors, errors


# 便捷函数
def convert_docx_to_html(docx_path: str) -> Dict[str, Any]:
    """将 .docx 转换为 HTML"""
    converter = DocxConverter()
    return converter.to_html(docx_path)


def convert_html_to_docx(
    html: str,
    output_path: str,
    original_docx_path: Optional[str] = None
) -> str:
    """将 HTML 转换为 .docx"""
    converter = DocxConverter()
    return converter.to_docx(html, output_path, original_docx_path)


def validate_docx_file(docx_path: str) -> tuple[bool, list[str]]:
    """验证 .docx 文件"""
    converter = DocxConverter()
    return converter.validate_docx(docx_path)
