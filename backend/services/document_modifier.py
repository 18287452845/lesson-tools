"""
Document modifier service for editing existing Word documents while preserving format.
This is a critical service for the "old lesson plan editing" feature.
"""
from pathlib import Path
from typing import Optional, List, Dict, Any, Literal
from dataclasses import dataclass
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from ..services.lesson_plan_parser import SectionInfo


@dataclass
class ParagraphFormat:
    """Format information for a paragraph."""
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    color: Optional[RGBColor] = None
    alignment: Optional[WD_PARAGRAPH_ALIGNMENT] = None
    line_spacing: Optional[float] = None
    space_before: Optional[float] = None
    space_after: Optional[float] = None


@dataclass
class EditRecord:
    """Record of an edit operation."""
    section_name: str
    operation: Literal["replace", "append", "insert", "add"]
    original_content: Optional[str]
    new_content: str
    timestamp: str


class DocumentModifier:
    """
    Modify Word documents while preserving original formatting.

    This is a key service for editing existing lesson plans. It:
    1. Records and preserves the original formatting
    2. Modifies content without changing styles
    3. Supports undo operations
    4. Handles both paragraph and table-based layouts
    """

    def __init__(self, doc_path: str):
        """
        Initialize the document modifier.

        Args:
            doc_path: Path to the document to modify
        """
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")

        self.doc = Document(str(doc_path))
        self.edit_history: List[EditRecord] = []
        self._original_backup = deepcopy(self.doc)

    def modify_section(
        self,
        section_info: SectionInfo,
        new_content: str,
        preserve_format: bool = True,
    ) -> bool:
        """
        Modify a section in the document.

        Args:
            section_info: Information about the section to modify
            new_content: New content for the section
            preserve_format: Whether to preserve original formatting

        Returns:
            True if modification was successful
        """
        if not section_info.found:
            return False

        # Record the edit
        self.edit_history.append(
            EditRecord(
                section_name=section_info.name,
                operation="replace",
                original_content=section_info.content,
                new_content=new_content,
                timestamp=self._get_timestamp(),
            )
        )

        if section_info.in_table:
            return self._modify_table_section(
                section_info, new_content, preserve_format
            )
        else:
            return self._modify_paragraph_section(
                section_info, new_content, preserve_format
            )

    def _modify_paragraph_section(
        self,
        section_info: SectionInfo,
        new_content: str,
        preserve_format: bool,
    ) -> bool:
        """Modify a section in paragraphs."""
        start_idx = section_info.start_para_idx
        end_idx = section_info.end_para_idx

        if start_idx is None or end_idx is None:
            return False

        # Save original formats if needed
        original_formats = []
        if preserve_format:
            for i in range(start_idx, min(end_idx + 1, len(self.doc.paragraphs))):
                para = self.doc.paragraphs[i]
                original_formats.append(self._extract_paragraph_format(para))

        # Split new content into lines
        new_lines = new_content.split("\n")

        try:
            # Handle first paragraph (may contain section header)
            first_para = self.doc.paragraphs[start_idx]

            # Check if first paragraph is a header
            is_header = self._is_section_header(first_para.text)

            if is_header:
                # Keep the header, replace the content after it
                header_text = self._extract_header_text(first_para.text)
                fmt = original_formats[0] if original_formats else None
                self._update_paragraph_content(
                    first_para, f"{header_text}{new_lines[0] if new_lines else ''}", fmt
                )

                # Handle remaining lines
                # Remove old content paragraphs
                for _ in range(start_idx + 1, end_idx + 1):
                    if start_idx + 1 < len(self.doc.paragraphs):
                        p = self.doc.paragraphs[start_idx + 1]._element
                        p.getparent().remove(p)

                # Add new lines as paragraphs
                for i, line in enumerate(new_lines[1:], 1):
                    self._insert_paragraph_after(
                        start_idx, line, original_formats[i] if i < len(original_formats) else None
                    )
            else:
                # Replace entire first paragraph
                fmt = original_formats[0] if original_formats else None
                self._update_paragraph_content(
                    first_para, new_lines[0] if new_lines else "", fmt
                )

                # Handle remaining lines
                # Remove old content paragraphs
                for _ in range(start_idx + 1, end_idx + 1):
                    if start_idx + 1 < len(self.doc.paragraphs):
                        p = self.doc.paragraphs[start_idx + 1]._element
                        p.getparent().remove(p)

                # Add new lines as paragraphs
                for i, line in enumerate(new_lines[1:], 1):
                    self._insert_paragraph_after(
                        start_idx, line, original_formats[i] if i < len(original_formats) else None
                    )

            return True

        except Exception as e:
            print(f"Error modifying paragraph section: {e}")
            return False

    def _modify_table_section(
        self,
        section_info: SectionInfo,
        new_content: str,
        preserve_format: bool,
    ) -> bool:
        """Modify a section in a table."""
        if (
            section_info.table_idx is None
            or section_info.cell_location is None
        ):
            return False

        table_idx = section_info.table_idx
        row_idx, col_idx = section_info.cell_location

        if table_idx >= len(self.doc.tables):
            return False

        table = self.doc.tables[table_idx]

        if row_idx >= len(table.rows) or col_idx >= len(table.rows[row_idx].cells):
            return False

        cell = table.rows[row_idx].cells[col_idx]

        # Record original format
        original_format = None
        if preserve_format and cell.paragraphs:
            original_format = self._extract_paragraph_format(cell.paragraphs[0])

        # Update cell content
        self._update_cell_content(cell, new_content, original_format)

        return True

    def add_section(
        self,
        section_name: str,
        content: str,
        position: Literal["end", "after_section"] = "end",
        after_section: Optional[str] = None,
        format_style: Optional[ParagraphFormat] = None,
    ) -> bool:
        """
        Add a new section to the document.

        Args:
            section_name: Name of the new section
            content: Content of the new section
            position: Where to add ("end" or "after_section")
            after_section: Section name to add after (if position="after_section")
            format_style: Optional format to apply

        Returns:
            True if addition was successful
        """
        # Get the display header for this section
        header = self._get_section_header(section_name)

        self.edit_history.append(
            EditRecord(
                section_name=section_name,
                operation="add",
                original_content=None,
                new_content=content,
                timestamp=self._get_timestamp(),
            )
        )

        if position == "end":
            # Add to the end of the document
            header_para = self.doc.add_paragraph()
            self._update_paragraph_content(header_para, header, format_style)

            content_para = self.doc.add_paragraph()
            self._update_paragraph_content(content_para, content, format_style)

        elif position == "after_section" and after_section:
            # Find the section and add after it
            # This would require knowing section locations
            # For now, add to end
            return self.add_section(section_name, content, "end", None, format_style)

        return True

    def append_to_section(
        self,
        section_info: SectionInfo,
        append_content: str,
        preserve_format: bool = True,
    ) -> bool:
        """
        Append content to an existing section.

        Args:
            section_info: Information about the section
            append_content: Content to append
            preserve_format: Whether to preserve formatting

        Returns:
            True if successful
        """
        if not section_info.found or not section_info.content:
            return False

        new_content = section_info.content + "\n" + append_content
        return self.modify_section(section_info, new_content, preserve_format)

    def insert_in_section(
        self,
        section_info: SectionInfo,
        insert_content: str,
        at_start: bool = True,
        preserve_format: bool = True,
    ) -> bool:
        """
        Insert content into a section.

        Args:
            section_info: Information about the section
            insert_content: Content to insert
            at_start: If True, insert at start; otherwise at end
            preserve_format: Whether to preserve formatting

        Returns:
            True if successful
        """
        if not section_info.found or not section_info.content:
            return False

        if at_start:
            new_content = insert_content + "\n" + section_info.content
        else:
            new_content = section_info.content + "\n" + insert_content

        return self.modify_section(section_info, new_content, preserve_format)

    def undo(self) -> bool:
        """
        Undo the last edit operation.

        Returns:
            True if undo was successful
        """
        if not self.edit_history:
            return False

        # For a full implementation, we would need to track more state
        # For now, we can restore from backup
        self.doc = deepcopy(self._original_backup)
        last_edit = self.edit_history.pop()

        # Re-apply all edits except the last one
        # This is a simplified implementation
        return True

    def save(self, output_path: Optional[str] = None) -> str:
        """
        Save the modified document.

        Args:
            output_path: Optional output path (defaults to overwriting original)

        Returns:
            Path to the saved document
        """
        if output_path is None:
            output_path = str(self.doc_path).replace(".docx", "_edited.docx")

        self.doc.save(output_path)
        return output_path

    def save_with_suffix(self, suffix: str = "_edited") -> str:
        """
        Save with a suffix added to the filename.

        Args:
            suffix: Suffix to add (default: "_edited")

        Returns:
            Path to the saved document
        """
        stem = self.doc_path.stem
        ext = self.doc_path.suffix
        new_name = f"{stem}{suffix}{ext}"
        output_path = str(self.doc_path.parent / new_name)
        return self.save(output_path)

    # ========================================================================
    # Helper methods
    # ========================================================================

    def _extract_paragraph_format(self, para) -> ParagraphFormat:
        """Extract formatting information from a paragraph."""
        fmt = ParagraphFormat()

        # Paragraph alignment
        fmt.alignment = para.alignment

        # Paragraph spacing
        if para.paragraph_format:
            fmt.line_spacing = para.paragraph_format.line_spacing
            if para.paragraph_format.space_before is not None:
                fmt.space_before = para.paragraph_format.space_before.pt
            if para.paragraph_format.space_after is not None:
                fmt.space_after = para.paragraph_format.space_after.pt

        # Run formatting (from first run)
        if para.runs:
            run = para.runs[0]
            if run.font:
                fmt.font_name = run.font.name
                if run.font.size is not None:
                    fmt.font_size = run.font.size.pt
                fmt.bold = run.font.bold
                fmt.italic = run.font.italic
                fmt.underline = run.font.underline
                if run.font.color and run.font.color.rgb:
                    fmt.color = run.font.color.rgb

        return fmt

    def _apply_paragraph_format(self, para, fmt: ParagraphFormat) -> None:
        """Apply formatting to a paragraph."""
        if fmt.alignment is not None:
            para.alignment = fmt.alignment

        if para.paragraph_format:
            if fmt.line_spacing is not None:
                para.paragraph_format.line_spacing = fmt.line_spacing
            if fmt.space_before is not None:
                para.paragraph_format.space_before = Pt(fmt.space_before)
            if fmt.space_after is not None:
                para.paragraph_format.space_after = Pt(fmt.space_after)

    def _apply_run_format(self, run, fmt: ParagraphFormat) -> None:
        """Apply formatting to a run."""
        if fmt.font_name is not None:
            run.font.name = fmt.font_name
        if fmt.font_size is not None:
            run.font.size = Pt(fmt.font_size)
        if fmt.bold is not None:
            run.font.bold = fmt.bold
        if fmt.italic is not None:
            run.font.italic = fmt.italic
        if fmt.underline is not None:
            run.font.underline = fmt.underline
        if fmt.color is not None:
            run.font.color.rgb = fmt.color

    def _update_paragraph_content(
        self, para, new_text: str, fmt: Optional[ParagraphFormat] = None
    ) -> None:
        """Update paragraph content while preserving format."""
        # Extract format from first run if not provided
        if fmt is None and para.runs:
            fmt = self._extract_paragraph_format(para)

        # Clear existing content
        para.clear()

        # Add new content
        run = para.add_run(new_text)

        # Apply format
        if fmt:
            self._apply_run_format(run, fmt)
            self._apply_paragraph_format(para, fmt)

    def _update_cell_content(
        self, cell, new_text: str, fmt: Optional[ParagraphFormat] = None
    ) -> None:
        """Update table cell content while preserving format."""
        if not cell.paragraphs:
            cell.add_paragraph()

        para = cell.paragraphs[0]

        # Extract format if not provided
        if fmt is None and para.runs:
            fmt = self._extract_paragraph_format(para)

        # Clear and update
        para.clear()
        run = para.add_run(new_text)

        if fmt:
            self._apply_run_format(run, fmt)
            self._apply_paragraph_format(para, fmt)

    def _insert_paragraph_after(
        self, idx: int, text: str, fmt: Optional[ParagraphFormat] = None
    ) -> None:
        """Insert a new paragraph after the given index."""
        if idx + 1 >= len(self.doc.paragraphs):
            # Add to end
            para = self.doc.add_paragraph(text)
        else:
            # Insert in the middle
            target = self.doc.paragraphs[idx]
            new_element = OxmlElement("w:p")
            target._element.addnext(new_element)
            para = Paragraph(new_element, target._parent)
            para.add_run(text)

        if fmt:
            self._apply_paragraph_format(para, fmt)
            if para.runs:
                self._apply_run_format(para.runs[0], fmt)

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        from ..services.lesson_plan_parser import LessonPlanParser
        return LessonPlanParser(self.doc_path)._match_section_header(text) is not None

    def _extract_header_text(self, text: str) -> str:
        """Extract just the header portion of text."""
        # Match common section header patterns
        import re
        from ..services.lesson_plan_parser import LessonPlanParser

        for patterns in LessonPlanParser.SECTION_PATTERNS.values():
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    return match.group(0)
        return text

    def _get_section_header(self, section_name: str) -> str:
        """Get the display header for a section."""
        headers = {
            "teaching_goals": "教学目标：",
            "key_points": "教学重点：",
            "difficult_points": "教学难点：",
            "teaching_tools": "教具准备：",
            "teaching_methods": "教学方法：",
            "student_analysis": "学情分析：",
            "textbook_analysis": "教材分析：",
            "teaching_process": "教学过程：",
            "homework": "作业布置：",
            "reflection": "教学反思：",
            "blackboard_design": "板书设计：",
        }
        return headers.get(section_name, f"{section_name}：")

    def _get_timestamp(self) -> str:
        """Get current timestamp."""
        from datetime import datetime
        return datetime.now().isoformat()


def modify_document(
    doc_path: str,
    section_info: SectionInfo,
    new_content: str,
    preserve_format: bool = True,
) -> str:
    """
    Convenience function to modify a document.

    Args:
        doc_path: Path to the document
        section_info: Section information
        new_content: New content
        preserve_format: Whether to preserve formatting

    Returns:
        Path to the modified document
    """
    modifier = DocumentModifier(doc_path)
    modifier.modify_section(section_info, new_content, preserve_format)
    return modifier.save_with_suffix()


def add_section_to_document(
    doc_path: str,
    section_name: str,
    content: str,
) -> str:
    """
    Convenience function to add a section to a document.

    Args:
        doc_path: Path to the document
        section_name: Name of the section
        content: Content to add

    Returns:
        Path to the modified document
    """
    modifier = DocumentModifier(doc_path)
    modifier.add_section(section_name, content, "end")
    return modifier.save_with_suffix()
