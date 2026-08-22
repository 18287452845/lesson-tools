"""Render fixed-format Yunlin teaching and experiment plans.

The two resources are immutable application assets. Generation copies the DOCX
package and changes only ``word/document.xml`` so styles, relationships, page
geometry, and all other package parts remain byte-for-byte identical.
"""

from __future__ import annotations

import hashlib
import re
import unicodedata
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Callable, Iterable, Literal, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from lxml import etree

from ..config import settings
from .batch_context import format_class_names
from .experiment_names import validate_experiment_chapters


CoursePlanType = Literal["teaching_plan", "experiment_plan"]

# 固定模板“作业”列每行最多 15 字；超长作业按内容回落为简短措辞，
# 兼顾评估实验与简单作业，避免整表都是同一句话。
HOMEWORK_MAX_CHARS = 15
HOMEWORK_BRIEF_RULES: tuple[tuple[tuple[str, ...], str], ...] = (
    (("实验报告", "实训报告", "报告"), "撰写实验报告"),
    (("练习", "习题", "题目"), "完成课后练习"),
    (("复习", "预习", "总结", "归纳"), "复习本课要点"),
    (("代码", "编程", "程序"), "完成编程练习"),
    (("实验", "实训", "上机", "实操"), "上机实操评估"),
    (("项目", "任务", "案例"), "完成项目任务"),
)
HOMEWORK_BRIEF_FALLBACK = "课后作业"
HOMEWORK_EMPTY_FALLBACK = "完成课后练习"

# 固定模板“教学重点/教学难点”列每行最多 25 字：
# 优先按分号保留完整要点句，装不下则只留前几条，单句超长才硬截断。
POINTS_MAX_CHARS = 25


def brief_point_line(line: str) -> str:
    """Condense focus/difficulty text to at most POINTS_MAX_CHARS per line.

    Multi-line values (e.g. AI-condensed points) are handled line by line;
    within one line whole semicolon clauses are kept when they fit.
    """
    brief_lines: list[str] = []
    for raw in str(line or "").splitlines():
        text = " ".join(raw.split())
        if not text:
            continue
        if len(text) <= POINTS_MAX_CHARS:
            brief_lines.append(text)
            continue
        kept: list[str] = []
        total = 0
        for part in re.split(r"[；;]", text):
            part = part.strip()
            if not part:
                continue
            extra = len(part) + (1 if kept else 0)
            if total + extra > POINTS_MAX_CHARS:
                break
            kept.append(part)
            total += extra
        brief_lines.append(
            "；".join(kept) if kept else text[:POINTS_MAX_CHARS]
        )
    return "\n".join(brief_lines)


def brief_homework_line(line: str) -> str:
    """Return the homework text as-is when short, else a matched brief phrase."""
    line = line.strip()
    if len(line) <= HOMEWORK_MAX_CHARS:
        return line
    for keywords, phrase in HOMEWORK_BRIEF_RULES:
        if any(keyword in line for keyword in keywords):
            return phrase
    return HOMEWORK_BRIEF_FALLBACK

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
W = f"{{{W_NS}}}"


@dataclass(frozen=True)
class CoursePlanTemplateSpec:
    type: CoursePlanType
    template_id: str
    name: str
    path: Path
    sha256: str
    title: str
    table_shapes: tuple[tuple[int, int], ...]
    capacity: int


TEMPLATE_SPECS: dict[CoursePlanType, CoursePlanTemplateSpec] = {
    "teaching_plan": CoursePlanTemplateSpec(
        type="teaching_plan",
        template_id="yunlin-teaching-plan",
        name="云林教师授课计划表",
        path=settings.teaching_plan_template_path,
        sha256="1e08531f22f93dc7cfa6a15d53a4f43829f1cb1feb8ce96aa1203920756607b8",
        title="云南林业职业技术学院教师授课计划表",
        table_shapes=((4, 9), (4, 9), (5, 9), (5, 9), (5, 9), (6, 9)),
        capacity=18,
    ),
    "experiment_plan": CoursePlanTemplateSpec(
        type="experiment_plan",
        template_id="yunlin-experiment-plan",
        name="云林课程实验计划表",
        path=settings.experiment_plan_template_path,
        sha256="76217eb3db84a0f1ca931381341665108d2c0135f36de692a543f1fbb812c5f1",
        title="云南林业职业技术学院课程实验计划表",
        table_shapes=((20, 5),),
        capacity=18,
    ),
}


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def validate_course_plan_template(plan_type: CoursePlanType) -> dict[str, Any]:
    """Validate the fixed asset checksum and the structural rendering contract."""
    spec = TEMPLATE_SPECS[plan_type]
    errors: list[str] = []
    warnings: list[str] = []
    actual_sha256 = ""

    if not spec.path.is_file():
        errors.append(f"固定模板不存在：{spec.path}")
    elif spec.path.suffix.lower() != ".docx":
        errors.append("固定模板必须是 .docx 文件")
    else:
        actual_sha256 = _file_sha256(spec.path)
        if actual_sha256 != spec.sha256:
            errors.append("固定模板指纹不匹配，资源可能已被替换或修改")

        try:
            document = Document(str(spec.path))
            if len(document.sections) != 1:
                errors.append(f"模板应为 1 节，实际为 {len(document.sections)} 节")
            elif document.sections[0].orientation != WD_ORIENT.LANDSCAPE:
                errors.append("模板页面方向必须为横向")

            table_shapes = tuple(
                (len(table.rows), len(table.columns)) for table in document.tables
            )
            if table_shapes != spec.table_shapes:
                errors.append(
                    f"模板表格结构不匹配：期望 {spec.table_shapes}，实际 {table_shapes}"
                )

            titles = [p.text.strip() for p in document.paragraphs]
            expected_title_count = 6 if plan_type == "teaching_plan" else 1
            if titles.count(spec.title) != expected_title_count:
                errors.append(
                    f"模板标题数量不匹配：期望 {expected_title_count}，"
                    f"实际 {titles.count(spec.title)}"
                )

            if plan_type == "teaching_plan":
                header = [cell.text.replace("\n", "") for cell in document.tables[0].rows[0].cells]
                for required in ("课序周次", "内容摘要", "教学时数", "教学重点", "教学难点"):
                    if not any(required.replace(" ", "") in value.replace(" ", "") for value in header):
                        errors.append(f"授课计划模板缺少表头：{required}")
            else:
                header = [cell.text.replace("\n", "") for cell in document.tables[0].rows[0].cells]
                for required in ("实验序号", "实验项目名称", "授课时间", "实验室", "备注"):
                    if required not in header:
                        errors.append(f"实验计划模板缺少表头：{required}")
        except Exception as exc:
            errors.append(f"固定模板无法解析：{exc}")

    if spec.path.is_file() and spec.path.stat().st_size < 10_000:
        warnings.append("模板文件体积异常偏小")

    return {
        "template_id": spec.template_id,
        "type": spec.type,
        "name": spec.name,
        "is_valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "file_size": spec.path.stat().st_size if spec.path.is_file() else 0,
        "sha256": actual_sha256,
        "expected_sha256": spec.sha256,
        "capacity": spec.capacity,
        "checked_at": datetime.now().astimezone().isoformat(),
    }


def validate_all_course_plan_templates() -> list[dict[str, Any]]:
    return [
        validate_course_plan_template("teaching_plan"),
        validate_course_plan_template("experiment_plan"),
    ]


def require_valid_course_plan_template(plan_type: CoursePlanType) -> dict[str, Any]:
    report = validate_course_plan_template(plan_type)
    if not report["is_valid"]:
        raise ValueError(
            f"{report['name']}校验失败：" + "；".join(report["errors"])
        )
    return report


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip(" .")
    return cleaned or "课程计划"


def _display_width(value: str) -> int:
    return sum(2 if unicodedata.east_asian_width(ch) in {"W", "F"} else 1 for ch in value)


def _fit_text(value: str, max_width: int) -> str:
    value = " ".join(str(value or "").split())
    if _display_width(value) <= max_width:
        return value
    result: list[str] = []
    width = 0
    for char in value:
        char_width = 2 if unicodedata.east_asian_width(char) in {"W", "F"} else 1
        if width + char_width + 2 > max_width:
            break
        result.append(char)
        width += char_width
    return "".join(result).rstrip() + "…"


def _normalize_chapters(chapters: Iterable[Any]) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for chapter in chapters:
        if hasattr(chapter, "model_dump"):
            item = chapter.model_dump()
        else:
            item = dict(chapter)
        item["topic"] = str(item.get("topic") or "").strip()
        item["content_summary"] = str(item.get("content_summary") or "").strip()
        item["key_points"] = str(item.get("key_points") or "").strip()
        item["difficult_points"] = str(item.get("difficult_points") or "").strip()
        concepts = item.get("key_concepts") or []
        item["key_concepts"] = [str(value).strip() for value in concepts if str(value).strip()]
        item["experiment_name"] = str(item.get("experiment_name") or "").strip()
        normalized.append(item)
    return normalized


def _group_chapters(chapters: Iterable[Any], size: int = 2) -> list[list[dict[str, Any]]]:
    values = _normalize_chapters(chapters)
    return [values[index:index + size] for index in range(0, len(values), size)]


def _direct_paragraphs(root: etree._Element) -> list[etree._Element]:
    return root.xpath("/w:document/w:body/w:p", namespaces=NS)


def _direct_tables(root: etree._Element) -> list[etree._Element]:
    return root.xpath("/w:document/w:body/w:tbl", namespaces=NS)


def _direct_runs(paragraph: etree._Element) -> list[etree._Element]:
    return paragraph.xpath("./w:r", namespaces=NS)


def _set_page_break_before(paragraph: etree._Element) -> None:
    paragraph_properties = paragraph.find(W + "pPr")
    if paragraph_properties is None:
        paragraph_properties = etree.Element(W + "pPr")
        paragraph.insert(0, paragraph_properties)
    if paragraph_properties.find(W + "pageBreakBefore") is None:
        etree.SubElement(paragraph_properties, W + "pageBreakBefore")


def _remove_page_break_before(paragraph: etree._Element) -> None:
    paragraph_properties = paragraph.find(W + "pPr")
    if paragraph_properties is None:
        return
    page_break = paragraph_properties.find(W + "pageBreakBefore")
    if page_break is not None:
        paragraph_properties.remove(page_break)


def _set_run_text(
    run: etree._Element,
    value: str,
    *,
    underline: bool | None = None,
) -> None:
    for child in list(run):
        if child.tag != W + "rPr":
            run.remove(child)
    if underline is not None:
        run_properties = run.find(W + "rPr")
        if run_properties is None:
            run_properties = etree.Element(W + "rPr")
            run.insert(0, run_properties)
        underline_element = run_properties.find(W + "u")
        if underline_element is None:
            underline_element = etree.SubElement(run_properties, W + "u")
        underline_element.set(W + "val", "single" if underline else "none")
    text = etree.SubElement(run, W + "t")
    if value.startswith(" ") or value.endswith(" "):
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value


def _set_cell_text(cell: etree._Element, value: str) -> None:
    paragraphs = cell.xpath("./w:p", namespaces=NS)
    if not paragraphs:
        paragraph = etree.SubElement(cell, W + "p")
    else:
        paragraph = paragraphs[0]
        for extra in paragraphs[1:]:
            cell.remove(extra)

    runs = paragraph.xpath("./w:r", namespaces=NS)
    if runs:
        run = runs[0]
    else:
        run = etree.SubElement(paragraph, W + "r")

    for child in list(paragraph):
        if child.tag not in {W + "pPr", W + "r"} or (
            child.tag == W + "r" and child is not run
        ):
            paragraph.remove(child)

    for child in list(run):
        if child.tag != W + "rPr":
            run.remove(child)

    lines = str(value or "").splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            etree.SubElement(run, W + "br")
        text = etree.SubElement(run, W + "t")
        if line.startswith(" ") or line.endswith(" "):
            text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text.text = line


def _set_cell_no_wrap(cell: etree._Element) -> None:
    """Keep validated short experiment names on one physical line."""
    cell_properties = cell.find(W + "tcPr")
    if cell_properties is None:
        cell_properties = etree.Element(W + "tcPr")
        cell.insert(0, cell_properties)
    if cell_properties.find(W + "noWrap") is None:
        etree.SubElement(cell_properties, W + "noWrap")


def _set_cell_runs(
    cell: etree._Element,
    segments: Sequence[tuple[str, bool]],
) -> None:
    """Replace cell text with explicitly styled runs while preserving layout."""
    paragraphs = cell.xpath("./w:p", namespaces=NS)
    if not paragraphs:
        paragraph = etree.SubElement(cell, W + "p")
    else:
        paragraph = paragraphs[0]
        for extra in paragraphs[1:]:
            cell.remove(extra)

    existing_runs = paragraph.xpath("./w:r", namespaces=NS)
    base_run = existing_runs[0] if existing_runs else etree.Element(W + "r")
    for child in list(paragraph):
        if child.tag != W + "pPr":
            paragraph.remove(child)

    for text, underline in segments:
        run = deepcopy(base_run)
        _set_run_text(run, text, underline=underline)
        paragraph.append(run)


def _set_experiment_schedule_cell(
    cell: etree._Element,
    *,
    week: int,
    weekday: str,
    periods: str,
    lesson_date: date,
) -> None:
    """Match Yunlin's schedule style with underlines on every variable value."""
    _set_cell_runs(
        cell,
        (
            ("第", False),
            (f" {week:02d} ", True),
            ("周  星期", False),
            (f" {weekday} ", True),
            ("  第", False),
            (f" {periods} ", True),
            ("节    ", False),
            (f" {lesson_date.month} ", True),
            ("月 ", False),
            (f" {lesson_date.day:02d} ", True),
            ("日", False),
        ),
    )


def _table_cells(row: etree._Element) -> list[etree._Element]:
    return row.xpath("./w:tc", namespaces=NS)


def _set_repeating_table_header(row: etree._Element, enabled: bool) -> None:
    row_properties = row.find(W + "trPr")
    if row_properties is None:
        row_properties = etree.Element(W + "trPr")
        row.insert(0, row_properties)
    for existing in list(row_properties.findall(W + "tblHeader")):
        row_properties.remove(existing)
    if enabled:
        etree.SubElement(row_properties, W + "tblHeader").set(W + "val", "true")


def _set_row_cant_split(row: etree._Element) -> None:
    row_properties = row.find(W + "trPr")
    if row_properties is None:
        row_properties = etree.Element(W + "trPr")
        row.insert(0, row_properties)
    if row_properties.find(W + "cantSplit") is None:
        etree.SubElement(row_properties, W + "cantSplit")


def _set_row_keep_with_next(row: etree._Element) -> None:
    """Keep a table row with the following row to avoid orphaned totals."""
    for paragraph in row.xpath("./w:tc/w:p", namespaces=NS):
        paragraph_properties = paragraph.find(W + "pPr")
        if paragraph_properties is None:
            paragraph_properties = etree.Element(W + "pPr")
            paragraph.insert(0, paragraph_properties)
        keep_next = paragraph_properties.find(W + "keepNext")
        if keep_next is None:
            keep_next = etree.SubElement(paragraph_properties, W + "keepNext")
        keep_next.set(W + "val", "true")


def _replace_run_with_field(
    run: etree._Element,
    instruction: str,
    display_value: str,
) -> None:
    parent = run.getparent()
    field = etree.Element(W + "fldSimple")
    field.set(W + "instr", f" {instruction} ")
    display_run = deepcopy(run)
    _set_run_text(display_run, display_value)
    field.append(display_run)
    parent.replace(run, field)


def _clear_cell_borders(cell: etree._Element) -> None:
    cell_properties = cell.find(W + "tcPr")
    if cell_properties is None:
        cell_properties = etree.Element(W + "tcPr")
        cell.insert(0, cell_properties)
    borders = cell_properties.find(W + "tcBorders")
    if borders is None:
        borders = etree.SubElement(cell_properties, W + "tcBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(W + edge_name)
        if edge is None:
            edge = etree.SubElement(borders, W + edge_name)
        edge.set(W + "val", "nil")


def _make_full_width_paragraph_row(
    template_row: etree._Element,
    paragraph: etree._Element,
    table: etree._Element,
) -> etree._Element:
    row = deepcopy(template_row)
    _set_repeating_table_header(row, False)
    _set_row_cant_split(row)
    cells = _table_cells(row)
    if len(cells) != 9:
        raise ValueError("授课计划签字行结构已改变")

    first_cell = cells[0]
    for cell in cells[1:]:
        row.remove(cell)
    for existing_paragraph in first_cell.xpath("./w:p", namespaces=NS):
        first_cell.remove(existing_paragraph)
    first_cell.append(deepcopy(paragraph))

    cell_properties = first_cell.find(W + "tcPr")
    if cell_properties is None:
        cell_properties = etree.Element(W + "tcPr")
        first_cell.insert(0, cell_properties)
    grid_span = cell_properties.find(W + "gridSpan")
    if grid_span is None:
        grid_span = etree.SubElement(cell_properties, W + "gridSpan")
    grid_span.set(W + "val", "9")
    widths = [
        int(column.get(W + "w"))
        for column in table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
        if column.get(W + "w")
    ]
    cell_width = cell_properties.find(W + "tcW")
    if cell_width is not None and widths:
        cell_width.set(W + "w", str(sum(widths)))
        cell_width.set(W + "type", "dxa")
    _clear_cell_borders(first_cell)
    return row


def _patch_docx(
    template_path: Path,
    output_path: Path,
    patcher: Callable[[etree._Element], None],
) -> Path:
    with zipfile.ZipFile(template_path, "r") as source:
        document_xml = source.read("word/document.xml")
        root = etree.fromstring(document_xml)
        patcher(root)
        patched_xml = etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_path, "w") as target:
            for item in source.infolist():
                payload = patched_xml if item.filename == "word/document.xml" else source.read(item.filename)
                copied = deepcopy(item)
                target.writestr(copied, payload)
    return output_path


def _parse_date(value: str | date) -> date:
    if isinstance(value, date):
        return value
    try:
        return date.fromisoformat(value)
    except ValueError as exc:
        raise ValueError("首课日期必须是 YYYY-MM-DD 格式") from exc


def _format_classes(class_names: Sequence[str]) -> str:
    return format_class_names(class_names)


class CoursePlanRenderer:
    """Create fixed teaching/experiment plan documents from a batch schedule."""

    def __init__(self, output_dir: Path | None = None):
        self.output_dir = Path(output_dir or settings.output_dir)

    @staticmethod
    def _patch_metadata_paragraph(
        paragraph: etree._Element,
        *,
        course_name: str,
        academic_year: str,
        semester: int,
        class_display: str,
        teacher_name: str,
    ) -> None:
        """Fill the copied metadata paragraph and turn page digits into fields."""
        runs = _direct_runs(paragraph)
        if len(runs) < 14:
            raise ValueError("授课计划元数据段结构已改变")
        _set_run_text(runs[1], f" {course_name}")
        _set_run_text(runs[4], academic_year)
        _set_run_text(runs[6], str(semester))
        _set_run_text(runs[8], "")
        _set_run_text(runs[9], "")
        _set_run_text(runs[10], class_display)
        _set_run_text(runs[11], "")
        _set_run_text(runs[13], teacher_name)

        marker_index = next(
            (
                index
                for index, run in enumerate(runs)
                if "页第" in "".join(run.xpath(".//w:t/text()", namespaces=NS))
            ),
            None,
        )
        if marker_index is None:
            raise ValueError("授课计划页码段结构已改变")
        total_index = next(
            (
                index
                for index in range(marker_index - 1, -1, -1)
                if "".join(runs[index].xpath(".//w:t/text()", namespaces=NS)).strip().isdigit()
            ),
            None,
        )
        current_index = next(
            (
                index
                for index in range(marker_index + 1, len(runs))
                if "".join(runs[index].xpath(".//w:t/text()", namespaces=NS)).strip().isdigit()
            ),
            None,
        )
        if total_index is None or current_index is None:
            raise ValueError("授课计划页码段结构已改变")
        _replace_run_with_field(runs[total_index], "NUMPAGES", "1")
        _replace_run_with_field(runs[current_index], "PAGE", "1")

    @staticmethod
    def _prepare_adaptive_teaching_layout(
        root: etree._Element,
        group_count: int,
        *,
        course_name: str,
        academic_year: str,
        semester: int,
        class_display: str,
        teacher_name: str,
    ) -> tuple[list[etree._Element], etree._Element]:
        """Build one flowing table whose first four rows repeat on every page.

        The title and metadata paragraphs from the template are composed into
        the table body itself (as repeated header rows) instead of a Word
        header part, so the patched package only touches word/document.xml.
        """
        paragraphs = _direct_paragraphs(root)
        tables = _direct_tables(root)
        if len(paragraphs) != 19 or len(tables) != 6:
            raise ValueError("授课计划模板分页结构已改变")

        title_paragraph = deepcopy(paragraphs[0])
        metadata_paragraph = deepcopy(paragraphs[1])
        _remove_page_break_before(metadata_paragraph)
        self_cls = CoursePlanRenderer
        self_cls._patch_metadata_paragraph(
            metadata_paragraph,
            course_name=course_name,
            academic_year=academic_year,
            semester=semester,
            class_display=class_display,
            teacher_name=teacher_name,
        )

        continuous_table = deepcopy(tables[0])
        continuous_rows = continuous_table.xpath("./w:tr", namespaces=NS)
        if len(continuous_rows) < 3:
            raise ValueError("授课计划表头结构已改变")
        for row in continuous_rows[2:]:
            continuous_table.remove(row)
        for row in continuous_rows[:2]:
            _set_repeating_table_header(row, True)
            _set_row_cant_split(row)

        source_rows = tables[5].xpath("./w:tr", namespaces=NS)
        if len(source_rows) != 6:
            raise ValueError("授课计划末页表格结构已改变")
        data_template = source_rows[2]

        title_row = _make_full_width_paragraph_row(
            data_template, title_paragraph, continuous_table
        )
        metadata_row = _make_full_width_paragraph_row(
            data_template, metadata_paragraph, continuous_table
        )
        for row in (title_row, metadata_row):
            _set_repeating_table_header(row, True)
        continuous_table.insert(0, title_row)
        continuous_table.insert(1, metadata_row)

        data_rows: list[etree._Element] = []
        for _ in range(group_count):
            row = deepcopy(data_template)
            _set_repeating_table_header(row, False)
            _set_row_cant_split(row)
            continuous_table.append(row)
            data_rows.append(row)

        total_row = deepcopy(source_rows[-1])
        _set_repeating_table_header(total_row, False)
        _set_row_cant_split(total_row)
        _set_row_keep_with_next(total_row)
        continuous_table.append(total_row)
        continuous_table.append(
            _make_full_width_paragraph_row(
                data_template,
                paragraphs[18],
                continuous_table,
            )
        )

        body = root.find(W + "body")
        if body is None:
            raise ValueError("授课计划模板正文结构已改变")
        section_properties = body.find(W + "sectPr")
        if section_properties is None:
            raise ValueError("授课计划模板缺少页面设置")
        for child in list(body):
            if child is not section_properties:
                body.remove(child)
        section_properties.addprevious(continuous_table)
        return data_rows, total_row

    def render_teaching_plan(
        self,
        *,
        batch_task_id: str,
        course_name: str,
        grade: str,
        class_names: Sequence[str],
        academic_year: str,
        semester: int,
        teacher_name: str,
        total_hours: int,
        hours_per_lesson: int,
        start_week: int,
        chapters: Iterable[Any],
        location: str | None = None,
    ) -> str:
        require_valid_course_plan_template("teaching_plan")
        groups = _group_chapters(chapters)
        spec = TEMPLATE_SPECS["teaching_plan"]
        if len(groups) > spec.capacity:
            raise ValueError(f"授课计划最多容纳 {spec.capacity} 周，当前为 {len(groups)} 周")

        class_display = _format_classes(class_names)
        if not class_display:
            raise ValueError("同步生成授课计划时必须选择至少一个班级")

        output_name = _safe_filename(
            f"{teacher_name}-{grade}《{course_name}》授课计划表.docx"
        )
        output_path = self.output_dir / f"{batch_task_id}_{output_name}"

        def patch(root: etree._Element) -> None:
            data_rows, total_row = self._prepare_adaptive_teaching_layout(
                root,
                len(groups),
                course_name=course_name,
                academic_year=academic_year,
                semester=semester,
                class_display=class_display,
                teacher_name=teacher_name,
            )

            for slot_index, row in enumerate(data_rows):
                cells = _table_cells(row)
                if len(cells) != 9:
                    raise ValueError("授课计划数据行结构已改变")

                group = groups[slot_index]
                weekly_hours = len(group) * hours_per_lesson
                theory_hours = weekly_hours // 2
                practice_hours = weekly_hours - theory_hours
                topics = [item["topic"] for item in group if item["topic"]]
                concepts = [
                    concept for item in group for concept in item.get("key_concepts", [])
                ]
                summaries = [
                    item["content_summary"] for item in group if item.get("content_summary")
                ]

                focus_values = [
                    brief_point_line(item["key_points"])
                    for item in group
                    if item["key_points"]
                ]
                focus = "\n".join(dict.fromkeys(focus_values))
                if not focus:
                    focus = brief_point_line("、".join(dict.fromkeys(concepts)))
                if not focus:
                    focus = brief_point_line(
                        summaries[0] if summaries else f"掌握{'、'.join(topics)}"
                    )
                difficulty_values = [
                    brief_point_line(item["difficult_points"])
                    for item in group
                    if item["difficult_points"]
                ]
                difficulty = "\n".join(dict.fromkeys(difficulty_values))
                if not difficulty:
                    difficulty = brief_point_line("、".join(dict.fromkeys(concepts[-2:])))
                if not difficulty:
                    difficulty = brief_point_line(
                        summaries[-1] if summaries else f"综合运用{'、'.join(topics)}"
                    )
                assignment_lines: list[str] = []
                for item in group:
                    homework = item.get("homework")
                    if isinstance(homework, Mapping):
                        assignment_lines.extend(
                            str(homework.get(field_name) or "").strip()
                            for field_name in ("required", "optional")
                            if str(homework.get(field_name) or "").strip()
                        )
                    elif str(homework or "").strip():
                        assignment_lines.append(str(homework).strip())
                assignment_lines = [
                    brief_homework_line(line) for line in assignment_lines
                ]
                assignment_lines = list(dict.fromkeys(assignment_lines)) or [
                    HOMEWORK_EMPTY_FALLBACK
                ]

                values = (
                    str(start_week + slot_index),
                    "\n".join(topics),
                    str(theory_hours),
                    str(practice_hours),
                    str(weekly_hours),
                    focus,
                    difficulty,
                    "机房、计算机",
                    "\n".join(assignment_lines),
                )
                for cell, value in zip(cells, values):
                    _set_cell_text(cell, value)

            total_cells = _table_cells(total_row)
            theory_total = total_hours // 2
            practice_total = total_hours - theory_total
            for index, value in ((1, "总课时"), (2, str(theory_total)),
                                 (3, str(practice_total)), (4, str(total_hours))):
                _set_cell_text(total_cells[index], value)

        return str(_patch_docx(spec.path, output_path, patch))

    def render_experiment_plans(
        self,
        *,
        batch_task_id: str,
        course_name: str,
        grade: str,
        class_names: Sequence[str],
        academic_year: str,
        semester: int,
        teacher_name: str,
        plan_date: str | date,
        first_class_date: str | date,
        class_periods: str,
        hours_per_lesson: int,
        start_week: int,
        chapters: Iterable[Any],
        location: str,
        class_schedules: Sequence[Mapping[str, Any]] | None = None,
    ) -> list[str]:
        require_valid_course_plan_template("experiment_plan")
        names = [str(name).strip() for name in class_names if str(name).strip()]
        if not names:
            raise ValueError("同步生成实验计划时必须选择至少一个班级")

        groups = _group_chapters(chapters)
        projects_by_group = dict(
            validate_experiment_chapters(
                [item for group in groups for item in group],
                require_every_group=False,
            )
        )
        schedule: list[tuple[int, list[dict[str, Any]], str]] = []
        for group_index, group in enumerate(groups):
            project = projects_by_group.get(group_index + 1)
            if not project:
                continue
            schedule.append((group_index, group, project))

        spec = TEMPLATE_SPECS["experiment_plan"]
        if len(schedule) > spec.capacity:
            raise ValueError(f"实验计划最多容纳 {spec.capacity} 条，当前为 {len(schedule)} 条")

        signed_date = _parse_date(plan_date)
        weekday_names = "一二三四五六日"
        schedule_by_class = {
            str(item.get("class_name") or "").strip(): item
            for item in (class_schedules or [])
            if str(item.get("class_name") or "").strip()
        }
        outputs: list[str] = []

        for class_name in names:
            class_schedule = schedule_by_class.get(class_name)
            if class_schedule:
                first_date = _parse_date(str(class_schedule.get("first_class_date") or ""))
                periods = str(class_schedule.get("class_periods") or "").strip()
                periods = periods.replace("第", "").replace("节", "")
                classroom = str(class_schedule.get("classroom") or "").strip()
                weekday_number = int(class_schedule.get("weekday") or 0)
                if weekday_number < 1 or weekday_number > 7:
                    raise ValueError(f"{class_name}的星期设置无效")
                if first_date.isoweekday() != weekday_number:
                    raise ValueError(f"{class_name}的第一周日期与星期设置不一致")
                weekday = weekday_names[weekday_number - 1]
            else:
                first_date = _parse_date(first_class_date)
                periods = str(class_periods).strip().replace("第", "").replace("节", "")
                classroom = location
                weekday = weekday_names[first_date.weekday()]

            output_name = _safe_filename(
                f"{teacher_name}-{class_name}《{course_name}》课程实验计划表.docx"
            )
            output_path = self.output_dir / f"{batch_task_id}_{output_name}"

            def patch(root: etree._Element) -> None:
                paragraphs = _direct_paragraphs(root)
                runs = _direct_runs(paragraphs[1])
                if len(runs) < 23:
                    raise ValueError("实验计划元数据段结构已改变")
                _set_run_text(runs[1], f" {course_name}", underline=True)
                _set_run_text(runs[3], f"{academic_year} ", underline=True)
                _set_run_text(runs[5], str(semester), underline=True)
                _set_run_text(runs[7], f"{class_name} ", underline=True)
                for run_index in (8, 9, 10, 11, 12):
                    _set_run_text(runs[run_index], "")
                _set_run_text(runs[14], teacher_name, underline=True)
                _set_run_text(runs[16], str(signed_date.year), underline=True)
                _set_run_text(runs[17], "")
                _set_run_text(runs[19], str(signed_date.month), underline=True)
                _set_run_text(runs[21], f"{signed_date.day:02d}", underline=True)

                table = _direct_tables(root)[0]
                rows = table.xpath("./w:tr", namespaces=NS)
                for row_index in range(1, 19):
                    if row_index > len(schedule):
                        table.remove(rows[row_index])
                        continue

                    _set_row_cant_split(rows[row_index])
                    cells = _table_cells(rows[row_index])
                    group_index, _group, project = schedule[row_index - 1]
                    lesson_date = first_date + timedelta(days=7 * group_index)
                    week = start_week + group_index
                    values = (
                        f"实验{row_index}",
                        project,
                        _fit_text(classroom, 18),
                        "",
                    )
                    for cell, value in zip((cells[0], cells[1], cells[3], cells[4]), values):
                        _set_cell_text(cell, value)
                    _set_cell_no_wrap(cells[1])
                    _set_experiment_schedule_cell(
                        cells[2],
                        week=week,
                        weekday=weekday,
                        periods=periods,
                        lesson_date=lesson_date,
                    )

                _set_row_cant_split(rows[19])
                total_cells = _table_cells(rows[19])
                total_hours = len(schedule) * hours_per_lesson
                _set_cell_text(total_cells[1], f"   {total_hours}    学时")

            outputs.append(str(_patch_docx(spec.path, output_path, patch)))

        return outputs
